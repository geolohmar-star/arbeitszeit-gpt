import io
import json
import logging

from django.conf import settings as django_settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.csrf import csrf_exempt

from .forms import BriefvorgangForm
from .models import Briefvorlage, Briefvorgang

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Hilfsfunktionen: DOCX-Befuellung
# ---------------------------------------------------------------------------

def _fuelle_vorlage(vorlage_bytes: bytes, platzhalter: dict) -> bytes:
    """Ersetzt alle {{schluessel}}-Platzhalter in der DOCX-Vorlage.

    Verarbeitet sowohl normale Absaetze als auch Tabellenzellen.
    Zusammengesetzte Platzhalter (ueber mehrere Runs gesplittet) werden
    korrekt behandelt, indem der gesamte Absatz-Text rekonstruiert wird.
    """
    from docx import Document

    doc = Document(io.BytesIO(vorlage_bytes))

    def _ersetze_absatz(para):
        """Ersetzt Platzhalter in einem einzelnen Absatz."""
        full_text = "".join(run.text for run in para.runs)
        neuer_text = full_text
        for schluessel, wert in platzhalter.items():
            token = "{{" + schluessel + "}}"
            neuer_text = neuer_text.replace(token, str(wert) if wert else "")
        if neuer_text != full_text and para.runs:
            # Ersetzten Text in ersten Run schreiben, restliche Runs leeren
            para.runs[0].text = neuer_text
            for run in para.runs[1:]:
                run.text = ""

    # Normale Absaetze
    for para in doc.paragraphs:
        _ersetze_absatz(para)

    # Tabellenzellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _ersetze_absatz(para)

    # Kopf- und Fusszeilen aller Abschnitte
    for section in doc.sections:
        for para in section.footer.paragraphs:
            _ersetze_absatz(para)
        for para in section.header.paragraphs:
            _ersetze_absatz(para)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _datum_deutsch(datum) -> str:
    """Formatiert ein date-Objekt als deutsches Datum (z. B. 15. M\xe4rz 2026)."""
    monate = [
        "", "Januar", "Februar", "M\xe4rz", "April", "Mai", "Juni",
        "Juli", "August", "September", "Oktober", "November", "Dezember",
    ]
    return f"{datum.day}. {monate[datum.month]} {datum.year}"


def _erstelle_platzhalter(vorgang: Briefvorgang) -> dict:
    """Erstellt das Platzhalter-Dict aus einem Briefvorgang."""
    return {
        # Absender
        "absender_name":    vorgang.absender_name,
        "absender_strasse": vorgang.absender_strasse,
        "absender_ort":     vorgang.absender_ort,
        "absender_telefon": vorgang.absender_telefon,
        "absender_email":   vorgang.absender_email,
        # Empfaenger
        "empfaenger_name":    vorgang.empfaenger_name,
        "empfaenger_zusatz":  vorgang.empfaenger_zusatz,
        "empfaenger_strasse": vorgang.empfaenger_strasse,
        "empfaenger_plz_ort": vorgang.empfaenger_plz_ort,
        "empfaenger_land":    vorgang.empfaenger_land,
        # Brief
        "ort_datum":           f"{vorgang.ort}, {_datum_deutsch(vorgang.datum)}",
        "datum":               _datum_deutsch(vorgang.datum),
        "betreff":             vorgang.betreff,
        "anrede":              vorgang.anrede,
        "brieftext":           vorgang.brieftext,
        "grussformel":         vorgang.grussformel,
        "unterschrift_name":   vorgang.unterschrift_name,
        "unterschrift_titel":  vorgang.unterschrift_titel,
        # Fusszeile – aus der Vorlage
        "fusszeile_firmenname": vorgang.vorlage.fusszeile_firmenname,
        "fusszeile_telefon":    vorgang.vorlage.fusszeile_telefon,
        "fusszeile_telefax":    vorgang.vorlage.fusszeile_telefax,
        "fusszeile_email":      vorgang.vorlage.fusszeile_email,
        "fusszeile_internet":   vorgang.vorlage.fusszeile_internet,
    }


# ---------------------------------------------------------------------------
# OnlyOffice: JWT-Signierung
# ---------------------------------------------------------------------------

def _onlyoffice_jwt(payload: dict) -> str:
    """Signiert den OnlyOffice-Konfigurations-Payload als HS256-JWT."""
    import jwt
    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if not secret:
        return ""
    return jwt.encode(payload, secret, algorithm="HS256")


# ---------------------------------------------------------------------------
# Views: Briefvorlagen-Verwaltung
# ---------------------------------------------------------------------------

@login_required
def vorlage_liste(request):
    """Listet alle aktiven Briefvorlagen auf."""
    vorlagen = Briefvorlage.objects.filter(ist_aktiv=True)
    return render(request, "korrespondenz/vorlage_liste.html", {"vorlagen": vorlagen})


# ---------------------------------------------------------------------------
# Views: Briefvorgaenge
# ---------------------------------------------------------------------------

@login_required
def brief_liste(request):
    """Listet alle Briefvorgaenge des aktuellen Benutzers auf."""
    briefe = Briefvorgang.objects.filter(erstellt_von=request.user)
    return render(request, "korrespondenz/brief_liste.html", {"briefe": briefe})


@login_required
def brief_erstellen(request):
    """Erstellt einen neuen Briefvorgang aus einer Vorlage.

    GET:  Zeigt das Formular mit allen Feldern.
    POST: Validiert, erstellt den Briefvorgang, befuellt die DOCX-Vorlage
          und leitet zur Detailseite weiter.
    """
    # Vorlagen-Vorauswahl: GET-Parameter hat Vorrang, sonst Standard-Vorlage
    vorlage_pk = request.GET.get("vorlage")
    if not vorlage_pk:
        standard = Briefvorlage.objects.filter(ist_aktiv=True, ist_standard=True).first()
        if standard:
            vorlage_pk = standard.pk
    initial = {}

    # Absender-Vorbelegung: Vorlage-Defaults haben Vorrang,
    # dann Firma-Einstellungen – Unterschrift immer aus HR-Profil
    if vorlage_pk:
        initial["vorlage"] = vorlage_pk
        try:
            vorlage = Briefvorlage.objects.get(pk=vorlage_pk, ist_aktiv=True)
            if vorlage.default_absender_name:
                initial["absender_name"]    = vorlage.default_absender_name
            if vorlage.default_absender_strasse:
                initial["absender_strasse"] = vorlage.default_absender_strasse
            if vorlage.default_absender_ort:
                initial["absender_ort"]     = vorlage.default_absender_ort
            if vorlage.default_absender_telefon:
                initial["absender_telefon"] = vorlage.default_absender_telefon
            if vorlage.default_absender_email:
                initial["absender_email"]   = vorlage.default_absender_email
            if vorlage.default_ort:
                initial["ort"]              = vorlage.default_ort
            if vorlage.default_grussformel:
                initial["grussformel"]      = vorlage.default_grussformel
        except Briefvorlage.DoesNotExist:
            pass

    # Unterschrift immer aus HR-Profil des Users
    try:
        mitarbeiter = request.user.hr_mitarbeiter
        initial.setdefault("unterschrift_name", mitarbeiter.vollname)
        if mitarbeiter.stelle:
            initial.setdefault("unterschrift_titel", mitarbeiter.stelle.bezeichnung)
    except Exception:
        initial.setdefault("unterschrift_name", request.user.get_full_name() or request.user.username)

    form = BriefvorgangForm(request.POST or None, initial=initial)

    if request.method == "POST" and form.is_valid():
        vorgang = form.save(commit=False)
        vorgang.erstellt_von = request.user
        vorgang.save()

        # DOCX-Vorlage mit Formulardaten befuellen
        try:
            vorlage_bytes = bytes(vorgang.vorlage.inhalt)
            platzhalter = _erstelle_platzhalter(vorgang)
            ausgefuellt = _fuelle_vorlage(vorlage_bytes, platzhalter)
            Briefvorgang.objects.filter(pk=vorgang.pk).update(inhalt=ausgefuellt)
        except Exception as exc:
            logger.error("Vorlage befuellen fehlgeschlagen fuer Brief %s: %s", vorgang.pk, exc)
            messages.warning(request, "Vorlage konnte nicht befuellt werden. Bitte in OnlyOffice pruefen.")

        messages.success(request, "Brief erstellt.")
        return redirect("korrespondenz:brief_detail", pk=vorgang.pk)

    # Vorlagen-Defaults als JSON fuer JS-Auto-Fill mitgeben
    vorlagen_defaults = {}
    for v in Briefvorlage.objects.filter(ist_aktiv=True):
        vorlagen_defaults[v.pk] = {
            "absender_name":    v.default_absender_name,
            "absender_strasse": v.default_absender_strasse,
            "absender_ort":     v.default_absender_ort,
            "absender_telefon": v.default_absender_telefon,
            "absender_email":   v.default_absender_email,
            "ort":              v.default_ort,
            "grussformel":      v.default_grussformel,
        }

    return render(request, "korrespondenz/brief_erstellen.html", {
        "form": form,
        "vorlagen_defaults": vorlagen_defaults,
    })


@login_required
def vorlage_defaults(request, pk):
    """Liefert die Standard-Absender-Felder einer Briefvorlage als JSON.

    Wird per AJAX aufgerufen wenn der User die Vorlage wechselt.
    """
    vorlage = get_object_or_404(Briefvorlage, pk=pk, ist_aktiv=True)
    return JsonResponse({
        "absender_name":    vorlage.default_absender_name,
        "absender_strasse": vorlage.default_absender_strasse,
        "absender_ort":     vorlage.default_absender_ort,
        "absender_telefon": vorlage.default_absender_telefon,
        "absender_email":   vorlage.default_absender_email,
        "ort":              vorlage.default_ort,
        "grussformel":      vorlage.default_grussformel,
    })


@login_required
def brief_detail(request, pk):
    """Zeigt die Detailseite eines Briefvorgangs."""
    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)
    return render(request, "korrespondenz/brief_detail.html", {"brief": brief})


@login_required
def brief_status_aendern(request, pk):
    """Aendert den Status eines Briefvorgangs per POST."""
    if request.method != "POST":
        return redirect("korrespondenz:brief_detail", pk=pk)

    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)
    neuer_status = request.POST.get("status", "")
    erlaubte = [s for s, _ in Briefvorgang.STATUS_CHOICES]

    if neuer_status in erlaubte:
        brief.status = neuer_status
        brief.save(update_fields=["status"])
        messages.success(request, f"Status geaendert: {brief.get_status_display()}")
    else:
        messages.error(request, "Ungueltiger Status.")

    return redirect("korrespondenz:brief_detail", pk=pk)


@login_required
def brief_loeschen(request, pk):
    """Loescht einen Briefvorgang nach Bestaetigung per POST."""
    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)

    if request.method == "POST":
        brief.delete()
        messages.success(request, "Brief geloescht.")
        return redirect("korrespondenz:brief_liste")

    return render(request, "korrespondenz/brief_loeschen_bestaetigung.html", {"brief": brief})


@login_required
def brief_in_ablage_speichern(request, pk):
    """Speichert den Briefvorgang als DOCX in die persoenliche DMS-Ablage (POST).

    Erstellt ein neues Dokument mit ist_persoenlich=True.
    Ist der Brief bereits archiviert, wird er nicht doppelt angelegt.
    """
    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)

    if request.method != "POST":
        return redirect("korrespondenz:brief_detail", pk=pk)

    if not brief.inhalt:
        messages.error(request, "Kein Inhalt zum Archivieren vorhanden.")
        return redirect("korrespondenz:brief_detail", pk=pk)

    from dms.models import Dokument
    from dms.services import speichere_dokument, suchvektor_befuellen

    dateiname = f"Brief_{brief.datum}_{brief.betreff[:40]}.docx".replace(" ", "_")
    inhalt_bytes = bytes(brief.inhalt)

    dok = Dokument(
        titel=f"{brief.betreff} ({brief.datum})",
        klasse="offen",
        beschreibung=f"Archivierter Brief an {brief.empfaenger_name}",
        dateiname=dateiname,
        dateityp="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        groesse_bytes=len(inhalt_bytes),
        erstellt_von=request.user,
        ist_persoenlich=True,
    )
    speichere_dokument(dok, inhalt_bytes)
    dok.save()
    suchvektor_befuellen(dok)

    messages.success(request, f'Brief wurde in deine persoenliche Ablage archiviert.')
    return redirect("dms:meine_ablage")


@login_required
def brief_download(request, pk):
    """Laed den aktuellen DOCX-Inhalt des Briefvorgangs herunter."""
    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)

    if not brief.inhalt:
        messages.error(request, "Kein Inhalt vorhanden.")
        return redirect("korrespondenz:brief_detail", pk=pk)

    dateiname = f"Brief_{brief.datum}_{brief.betreff[:30]}.docx".replace(" ", "_")
    return HttpResponse(
        bytes(brief.inhalt),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{dateiname}"'},
    )


# ---------------------------------------------------------------------------
# Hilfsfunktion: DOCX → PDF via OnlyOffice ConvertService
# ---------------------------------------------------------------------------

def _konvertiere_docx_zu_pdf(brief: Briefvorgang):
    """Sendet den DOCX-Inhalt des Briefvorgangs an den OnlyOffice ConvertService.

    OnlyOffice konvertiert server-seitig zu PDF und gibt eine Download-URL zurueck.
    Gibt die PDF-Bytes zurueck oder None bei Fehler.
    """
    import hashlib
    import time
    import urllib.request as urlreq

    onlyoffice_internal = (
        getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "").rstrip("/")
        or getattr(django_settings, "ONLYOFFICE_URL", "").rstrip("/")
    )
    prima_base = (
        getattr(django_settings, "PRIMA_ONLYOFFICE_BASE_URL", "").rstrip("/")
        or getattr(django_settings, "PRIMA_BASE_URL", "").rstrip("/")
    )
    if not onlyoffice_internal or not prima_base:
        logger.error("OnlyOffice oder PRIMA-Base-URL nicht konfiguriert.")
        return None

    # Einzigartiger Schluessel pro Konvertierungsauftrag
    key = hashlib.md5(
        f"convert-{brief.pk}-{time.time()}".encode()
    ).hexdigest()[:20]

    nutzlast = {
        "async":      False,
        "filetype":   "docx",
        "key":        key,
        "outputtype": "pdf",
        "title":      f"Brief_{brief.pk}.docx",
        "url":        f"{prima_base}/korrespondenz/{brief.pk}/onlyoffice/laden/",
    }
    nutzlast_bytes = json.dumps(nutzlast).encode("utf-8")

    headers = {"Content-Type": "application/json"}
    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if secret:
        import jwt as pyjwt
        token = pyjwt.encode(nutzlast, secret, algorithm="HS256")
        headers["Authorization"] = f"Bearer {token}"

    convert_url = f"{onlyoffice_internal}/ConvertService.ashx"
    try:
        req = urlreq.Request(convert_url, data=nutzlast_bytes, headers=headers, method="POST")
        with urlreq.urlopen(req, timeout=60) as resp:
            antwort_bytes = resp.read()
    except Exception as exc:
        logger.error("OnlyOffice Konvertierung fehlgeschlagen fuer Brief %s: %s", brief.pk, exc)
        return None

    # ConvertService antwortet mit XML (nicht JSON)
    # Beispiel: <FileResult><FileUrl>...</FileUrl><EndConvert>True</EndConvert></FileResult>
    import xml.etree.ElementTree as ET
    try:
        root = ET.fromstring(antwort_bytes)
        end_convert = (root.findtext("EndConvert") or "").strip().lower() == "true"
        pdf_url = (root.findtext("FileUrl") or "").strip()
        fehler_code = root.findtext("Error")
    except ET.ParseError as exc:
        logger.error("OnlyOffice XML-Antwort nicht parsebar fuer Brief %s: %s | Antwort: %s",
                     brief.pk, exc, antwort_bytes[:200])
        return None

    if fehler_code:
        logger.error("OnlyOffice Konvertierungsfehler fuer Brief %s: Fehlercode %s", brief.pk, fehler_code)
        return None

    if not end_convert:
        logger.error("OnlyOffice Konvertierung unvollstaendig fuer Brief %s", brief.pk)
        return None

    if not pdf_url:
        return None

    try:
        with urlreq.urlopen(pdf_url, timeout=30) as resp:
            return resp.read()
    except Exception as exc:
        logger.error("PDF-Download von OnlyOffice fehlgeschlagen fuer Brief %s: %s", brief.pk, exc)
        return None


def _erstelle_signaturseite(brief: Briefvorgang, request) -> bytes:
    """Erzeugt eine professionelle Signaturseite als PDF via WeasyPrint.

    Die Seite enthaelt Dokumentinfo, Signaturtyp-Erklaerung und einen
    klar definierten Stempelbereich fuer pyhanko und sign-me.

    Koordinaten des Stempelbereichs auf der PDF-Seite (A4 = 595x842pt):
      pyhanko box:            (20, 539, 502, 667)
      sign-me visibleSignature: x=20, y=175, width=482, height=128
    """
    from weasyprint import HTML
    from django.template.loader import render_to_string
    from django.utils import timezone
    from django.conf import settings as conf

    # Unterzeichner-Infos aus HR-Profil
    unterzeichner_name = request.user.get_full_name() or request.user.username
    unterzeichner_funktion = ""
    try:
        hr = request.user.hr_mitarbeiter
        if hr.stelle:
            unterzeichner_funktion = hr.stelle.bezeichnung
    except Exception:
        pass

    backend_name = getattr(conf, "SIGNATUR_BACKEND", "intern")
    signatur_backend = "PRIMA FES (intern)" if backend_name == "intern" else "Bundesdruckerei sign-me (QES)"

    kontext = {
        "brief": brief,
        "unterzeichner_name": unterzeichner_name,
        "unterzeichner_funktion": unterzeichner_funktion,
        "signiert_am": timezone.localtime(timezone.now()).strftime("%d.%m.%Y %H:%M Uhr"),
        "signatur_backend": signatur_backend,
    }
    html_string = render_to_string(
        "korrespondenz/signaturseite_pdf.html", kontext, request=request
    )
    return HTML(
        string=html_string, base_url=request.build_absolute_uri("/")
    ).write_pdf()


def _pdfs_zusammenfuehren(pdf_a: bytes, pdf_b: bytes) -> bytes:
    """Haengt pdf_b seitenweise an pdf_a an und gibt das Ergebnis zurueck."""
    import io
    from pypdf import PdfWriter, PdfReader

    writer = PdfWriter()
    for page in PdfReader(io.BytesIO(pdf_a)).pages:
        writer.add_page(page)
    for page in PdfReader(io.BytesIO(pdf_b)).pages:
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


@login_required
def brief_pdf_exportieren(request, pk):
    """Konvertiert den Brief zu PDF, haengt eine Signaturseite an und signiert digital.

    Ablauf:
      1. DOCX → PDF via OnlyOffice ConvertService
      2. Signaturseite als separates PDF erzeugen (WeasyPrint)
      3. Brief-PDF + Signaturseite zusammenfuehren (pypdf)
      4. Gesamtes PDF digital signieren via signatur.services.signiere_pdf()
         Der Stempel landet auf der Signaturseite (letzte Seite) im
         definierten Stempelbereich:
           pyhanko: box=(20, 539, 502, 667)
           sign-me: x=20, y=175, width=482, height=128
      5. Signiertes PDF als Download zurueckgeben
    """
    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)

    if not brief.inhalt:
        messages.error(request, "Kein Inhalt vorhanden.")
        return redirect("korrespondenz:brief_detail", pk=pk)

    # Schritt 1: DOCX → PDF via OnlyOffice
    brief_pdf = _konvertiere_docx_zu_pdf(brief)
    if not brief_pdf:
        messages.error(request, "PDF-Konvertierung fehlgeschlagen. OnlyOffice erreichbar?")
        return redirect("korrespondenz:brief_detail", pk=pk)

    # Schritt 2: Signaturseite erzeugen
    try:
        sig_seite = _erstelle_signaturseite(brief, request)
    except Exception as exc:
        logger.warning("Signaturseite konnte nicht erzeugt werden fuer Brief %s: %s", pk, exc)
        sig_seite = None

    # Schritt 3: Zusammenfuehren
    if sig_seite:
        try:
            pdf_bytes = _pdfs_zusammenfuehren(brief_pdf, sig_seite)
        except Exception as exc:
            logger.warning("PDF-Merge fehlgeschlagen fuer Brief %s: %s – nur Brief-PDF wird verwendet", pk, exc)
            pdf_bytes = brief_pdf
    else:
        pdf_bytes = brief_pdf

    # Schritt 4: Signieren
    # Der Stempel landet automatisch auf der letzten Seite (Signaturseite) im
    # Bereich (20, 539, 502, 667) – passt exakt zum visuellen Stempelrahmen.
    # Fuer sign-me: SIGNATUR_SIGN_ME_URL konfigurieren und visibleSignature-
    # Koordinaten (x=20, y=175, width=482, height=128) im Backend setzen.
    dateiname = f"Brief_{brief.datum}_{brief.betreff[:40]}.pdf".replace(" ", "_")
    try:
        from signatur.services import signiere_pdf
        pdf_bytes = signiere_pdf(
            pdf_bytes,
            request.user,
            dokument_name=dateiname,
            seite=-1,  # immer letzte Seite = Signaturseite
        )
    except Exception as exc:
        logger.warning(
            "PDF-Signierung fehlgeschlagen fuer Brief %s: %s – unsigniertes PDF wird ausgegeben",
            pk, exc,
        )

    # Schritt 5: Download
    return HttpResponse(
        pdf_bytes,
        content_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{dateiname}"'},
    )


# ---------------------------------------------------------------------------
# Views: OnlyOffice-Integration
# ---------------------------------------------------------------------------

@login_required
def brief_editor(request, pk):
    """Oeffnet das OnlyOffice-Editorfenster fuer einen Briefvorgang.

    Generiert eine JWT-signierte Editor-Konfiguration und gibt die
    Editor-HTML-Seite zurueck.
    """
    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)

    if not brief.inhalt:
        messages.error(request, "Kein Inhalt vorhanden. Bitte Brief neu erstellen.")
        return redirect("korrespondenz:brief_detail", pk=pk)

    onlyoffice_url = getattr(django_settings, "ONLYOFFICE_URL", "")
    if not onlyoffice_url:
        messages.error(request, "OnlyOffice ist nicht konfiguriert.")
        return redirect("korrespondenz:brief_detail", pk=pk)

    prima_base = (
        getattr(django_settings, "PRIMA_ONLYOFFICE_BASE_URL", "").rstrip("/")
        or getattr(django_settings, "PRIMA_BASE_URL", "").rstrip("/")
    )

    # Eindeutiger Cache-Key pro Brief + Version
    doc_key = f"korrespondenz-{brief.pk}-v{brief.version}"

    config = {
        "document": {
            "fileType": "docx",
            "key":      doc_key,
            "title":    f"Brief_{brief.datum}_{brief.betreff[:40]}.docx",
            "url":      f"{prima_base}/korrespondenz/{brief.pk}/onlyoffice/laden/",
        },
        "documentType": "word",
        "editorConfig": {
            "callbackUrl": f"{prima_base}/korrespondenz/{brief.pk}/onlyoffice/callback/",
            "lang":        "de-DE",
            "mode":        "edit",
            "user": {
                "id":   str(request.user.pk),
                "name": request.user.get_full_name() or request.user.username,
            },
            "customization": {
                "spellcheck": True,
            },
        },
    }

    token = _onlyoffice_jwt(config)

    return render(request, "korrespondenz/onlyoffice_editor.html", {
        "brief":          brief,
        "onlyoffice_url": onlyoffice_url,
        "config":         config,   # dict – json_script serialisiert direkt
        "token":          token,
    })


def brief_onlyoffice_laden(request, pk):
    """Liefert den DOCX-Inhalt des Briefvorgangs an den OnlyOffice-Server.

    Kein login_required – Authentifizierung erfolgt per JWT-Header.
    """
    import jwt

    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if secret:
        auth_header = request.headers.get("Authorization", "")
        if not auth_header.startswith("Bearer "):
            return HttpResponse("Unauthorized", status=401)
        try:
            jwt.decode(auth_header[7:], secret, algorithms=["HS256"])
        except jwt.PyJWTError:
            return HttpResponse("Unauthorized", status=401)

    brief = get_object_or_404(Briefvorgang, pk=pk)

    if not brief.inhalt:
        return HttpResponse("Kein Inhalt", status=404)

    return HttpResponse(
        bytes(brief.inhalt),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@csrf_exempt
def brief_onlyoffice_callback(request, pk):
    """Empfaengt die bearbeitete Datei vom OnlyOffice-Server und speichert sie.

    Status 2: Dokument bereit zum Speichern (alle Editoren geschlossen).
    Status 6: Forciert gespeichert.
    """
    import urllib.request
    import jwt

    if request.method != "POST":
        return JsonResponse({"error": 0})

    try:
        daten = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"error": 1})

    # JWT-Verifizierung
    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if secret:
        auth_header = request.headers.get("Authorization", "")
        if auth_header.startswith("Bearer "):
            try:
                jwt.decode(auth_header[7:], secret, algorithms=["HS256"])
            except jwt.PyJWTError:
                logger.warning("Korrespondenz Callback: ungueltige JWT fuer Brief %s", pk)
                return JsonResponse({"error": 1})

    status = daten.get("status")
    if status not in (2, 6):
        return JsonResponse({"error": 0})

    download_url = daten.get("url")
    if not download_url:
        return JsonResponse({"error": 0})

    # Oeffentliche URL auf interne umschreiben (Cloudflare entfernt Authorization-Header)
    oo_public   = getattr(django_settings, "ONLYOFFICE_URL", "").rstrip("/")
    oo_internal = getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "").rstrip("/")
    if oo_public and oo_internal and download_url.startswith(oo_public):
        download_url = download_url.replace(oo_public, oo_internal, 1)

    brief = get_object_or_404(Briefvorgang, pk=pk)

    try:
        dl_req = urllib.request.Request(download_url)
        if secret:
            dl_token = jwt.encode({"url": download_url}, secret, algorithm="HS256")
            dl_req.add_header("Authorization", f"Bearer {dl_token}")
        with urllib.request.urlopen(dl_req, timeout=30) as resp:
            neuer_inhalt = resp.read()
    except Exception as exc:
        logger.error("Korrespondenz Callback: Download fehlgeschlagen fuer Brief %s: %s | url=%s", pk, exc, download_url)
        return JsonResponse({"error": 1})

    neue_version = brief.version + 1
    Briefvorgang.objects.filter(pk=pk).update(
        inhalt=neuer_inhalt,
        version=neue_version,
    )

    logger.info("Korrespondenz Callback: Brief %s als Version %s gespeichert", pk, neue_version)
    return JsonResponse({"error": 0})


@login_required
def brief_onlyoffice_version_check(request, pk):
    """Gibt die aktuelle Versionsnummer des Briefvorgangs zurueck (fuer Polling)."""
    brief = get_object_or_404(Briefvorgang, pk=pk)
    return JsonResponse({"version": brief.version})


@login_required
def brief_onlyoffice_forcesave(request, pk):
    """Loest einen Force-Save im OnlyOffice Command Service aus."""
    import urllib.request as urlreq
    import urllib.parse

    if request.method != "POST":
        return JsonResponse({"ok": False, "fehler": "Nur POST erlaubt"})

    brief = get_object_or_404(Briefvorgang, pk=pk, erstellt_von=request.user)

    onlyoffice_internal = (
        getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "").rstrip("/")
        or getattr(django_settings, "ONLYOFFICE_URL", "").rstrip("/")
    )
    if not onlyoffice_internal:
        return JsonResponse({"ok": False, "fehler": "OnlyOffice nicht konfiguriert"})

    doc_key = f"korrespondenz-{brief.pk}-v{brief.version}"
    command_url = f"{onlyoffice_internal}/coauthoring/CommandService.ashx"

    payload = json.dumps({"c": "forcesave", "key": doc_key}).encode("utf-8")
    headers = {"Content-Type": "application/json"}

    # JWT-Header beifuegen wenn Secret konfiguriert
    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if secret:
        cmd_token = _onlyoffice_jwt({"c": "forcesave", "key": doc_key})
        headers["Authorization"] = f"Bearer {cmd_token}"

    try:
        req = urlreq.Request(command_url, data=payload, headers=headers, method="POST")
        with urlreq.urlopen(req, timeout=10) as resp:
            antwort = json.loads(resp.read())
        if antwort.get("error", 1) not in (0, 4):
            logger.warning("OnlyOffice CommandService Fehler fuer Brief %s: %s", pk, antwort)
            return JsonResponse({"ok": False, "fehler": str(antwort)})
    except Exception as exc:
        logger.error("OnlyOffice CommandService Aufruf fehlgeschlagen fuer Brief %s: %s", pk, exc)
        return JsonResponse({"ok": False, "fehler": str(exc)})

    return JsonResponse({"ok": True})
