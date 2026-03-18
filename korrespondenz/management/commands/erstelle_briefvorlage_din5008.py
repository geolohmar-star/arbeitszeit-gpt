"""Management-Command: Legt eine Standard-Briefvorlage nach DIN 5008 an.

Verwendung:
    python manage.py erstelle_briefvorlage_din5008

Die DOCX-Datei wird mit python-docx erstellt und als Briefvorlage-Objekt
in der Datenbank gespeichert. Platzhalter haben die Form {{schluessel}}.

Unterstuetzte Platzhalter:
    {{absender_name}}       Name des Absenders / der Firma
    {{absender_strasse}}    Strasse und Hausnummer des Absenders
    {{absender_ort}}        PLZ und Ort des Absenders
    {{absender_telefon}}    Telefonnummer des Absenders
    {{absender_email}}      E-Mail-Adresse des Absenders
    {{empfaenger_name}}     Name des Empfaengers
    {{empfaenger_zusatz}}   Abteilung / Zusatzinformation
    {{empfaenger_strasse}}  Strasse und Hausnummer des Empfaengers
    {{empfaenger_plz_ort}}  PLZ und Ort des Empfaengers
    {{empfaenger_land}}     Land des Empfaengers (leer = Deutschland)
    {{ort_datum}}           Ort und Datum als Zeichenkette (z. B. Musterstadt, 15. Maerz 2026)
    {{betreff}}             Betreff des Briefes
    {{anrede}}              Anredezeile (z. B. Sehr geehrte Frau Mueller,)
    {{brieftext}}           Textkörper des Briefes
    {{grussformel}}         Abschlussformel (z. B. Mit freundlichen Gruessen)
    {{unterschrift_name}}   Name des Unterzeichners
    {{unterschrift_titel}}  Funktion / Titel des Unterzeichners
"""
import io

from django.core.management.base import BaseCommand


class Command(BaseCommand):
    help = "Erstellt eine Standard-Briefvorlage nach DIN 5008 in der Datenbank."

    def add_arguments(self, parser):
        parser.add_argument(
            "--titel",
            default="Geschaeftsbrief DIN 5008",
            help="Titel der Briefvorlage (Standard: 'Geschaeftsbrief DIN 5008')",
        )
        parser.add_argument(
            "--ueberschreiben",
            action="store_true",
            help="Vorhandene Vorlage mit gleichem Titel ueberschreiben",
        )

    def handle(self, *args, **options):
        from korrespondenz.models import Briefvorlage

        titel = options["titel"]

        # DOCX-Vorlage erstellen
        docx_bytes = self._erstelle_din5008_docx()

        vorlage_qs = Briefvorlage.objects.filter(titel=titel)
        if vorlage_qs.exists():
            if not options["ueberschreiben"]:
                self.stdout.write(self.style.WARNING(
                    f"Vorlage '{titel}' existiert bereits. "
                    f"Mit --ueberschreiben aktualisieren."
                ))
                return
            # Nur das DOCX-Inhalt-Feld aktualisieren (keine Loeschung wegen FK-Schutz)
            vorlage_qs.update(inhalt=docx_bytes)
            self.stdout.write(self.style.SUCCESS(f"Briefvorlage '{titel}' erfolgreich aktualisiert."))
        else:
            Briefvorlage.objects.create(
                titel=titel,
                beschreibung=(
                    "Standard-Geschaeftsbrief nach DIN 5008 Typ A. "
                    "Enthaelt alle Standardfelder fuer die externe Kommunikation."
                ),
                inhalt=docx_bytes,
                ist_aktiv=True,
            )
            self.stdout.write(self.style.SUCCESS(f"Briefvorlage '{titel}' erfolgreich angelegt."))

    # ---------------------------------------------------------------------------
    # DOCX-Erzeugung
    # ---------------------------------------------------------------------------

    def _erstelle_din5008_docx(self) -> bytes:
        """Erzeugt ein DIN-5008-konformes DOCX mit Platzhaltern."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor

        doc = Document()

        # Seitenraender nach DIN 5008 Typ A
        section = doc.sections[0]
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.7)
        section.bottom_margin = Cm(2.0)

        # Standardschrift auf Arial 11pt setzen
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after  = Pt(0)

        # Dokumentsprache auf Deutsch setzen (Rechtschreibpruefung in OnlyOffice)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        rpr = style.element.get_or_add_rPr()
        for vorhandenes in rpr.findall(_qn("w:lang")):
            rpr.remove(vorhandenes)
        lang_el = OxmlElement("w:lang")
        lang_el.set(_qn("w:val"), "de-DE")
        lang_el.set(_qn("w:eastAsia"), "zh-CN")
        lang_el.set(_qn("w:bidi"), "ar-SA")
        rpr.append(lang_el)
        settings_elem = doc.settings.element
        for vorhandenes in settings_elem.findall(_qn("w:themeFontLang")):
            settings_elem.remove(vorhandenes)
        theme_lang = OxmlElement("w:themeFontLang")
        theme_lang.set(_qn("w:val"), "de-DE")
        settings_elem.append(theme_lang)

        def absatz(text="", fett=False, kursiv=False, schriftgroesse=11,
                   abstand_nach=0, abstand_vor=0, ausrichtung=None) -> None:
            """Fuegt einen Absatz mit einem einzelnen Run hinzu."""
            p = doc.add_paragraph()
            if text:
                run = p.add_run(text)
                run.font.name  = "Arial"
                run.font.size  = Pt(schriftgroesse)
                run.font.bold  = fett
                run.font.italic = kursiv
            pf = p.paragraph_format
            pf.space_before = Pt(abstand_vor)
            pf.space_after  = Pt(abstand_nach)
            if ausrichtung:
                p.alignment = ausrichtung

        # ---- Absenderkurzangabe (8 pt, grau) ----
        # Wird oberhalb des Anschriftenfeldes als eine Zeile dargestellt
        p_abs = doc.add_paragraph()
        run = p_abs.add_run(
            "{{absender_name}}  \u00b7  {{absender_strasse}}  \u00b7  {{absender_ort}}"
        )
        run.font.name  = "Arial"
        run.font.size  = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p_abs.paragraph_format.space_before = Pt(0)
        p_abs.paragraph_format.space_after  = Pt(4)

        # Trennlinie unter Absenderkurzangabe
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p_abs._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

        # ---- Leerzeile ----
        absatz()

        # ---- Empfaengeranschrift ----
        absatz("{{empfaenger_name}}", abstand_nach=0)
        absatz("{{empfaenger_zusatz}}", abstand_nach=0)
        absatz("{{empfaenger_strasse}}", abstand_nach=0)
        absatz("{{empfaenger_plz_ort}}", abstand_nach=0)
        absatz("{{empfaenger_land}}", abstand_nach=0)

        # ---- Leerzeilen (DIN 5008: mindestens 3 Leerzeilen nach Anschrift) ----
        absatz(abstand_nach=4)
        absatz(abstand_nach=4)
        absatz(abstand_nach=4)

        # ---- Ort und Datum (rechtsbündig) ----
        absatz(
            "{{ort_datum}}",
            ausrichtung=WD_ALIGN_PARAGRAPH.RIGHT,
            abstand_nach=6,
        )

        # ---- Betreff (fett) ----
        absatz("{{betreff}}", fett=True, abstand_nach=12)

        # ---- Anrede ----
        absatz("{{anrede}}", abstand_nach=12)

        # ---- Brieftext ----
        absatz("{{brieftext}}", abstand_nach=24)

        # ---- Grussformel ----
        absatz("{{grussformel}}", abstand_nach=0)

        # ---- Platz fuer Unterschrift (3 Leerzeilen) ----
        absatz(abstand_nach=4)
        absatz(abstand_nach=4)
        absatz(abstand_nach=4)

        # ---- Unterzeichner ----
        absatz("{{unterschrift_name}}", fett=True, abstand_nach=0)
        absatz("{{unterschrift_titel}}", schriftgroesse=9, abstand_nach=0)

        # ---- Echte DOCX-Fusszeile (erscheint auf jeder Seite unten) ----
        footer = section.footer
        footer.is_linked_to_previous = False

        # Vorhandenen leeren Absatz der Fusszeile nutzen
        if footer.paragraphs:
            p_fuss = footer.paragraphs[0]
        else:
            p_fuss = footer.add_paragraph()

        # Trennlinie oben an der Fusszeile
        pPr_fuss = p_fuss._p.get_or_add_pPr()
        pBdr_fuss = OxmlElement("w:pBdr")
        top_fuss = OxmlElement("w:top")
        top_fuss.set(qn("w:val"), "single")
        top_fuss.set(qn("w:sz"), "4")
        top_fuss.set(qn("w:space"), "1")
        top_fuss.set(qn("w:color"), "AAAAAA")
        pBdr_fuss.append(top_fuss)
        pPr_fuss.append(pBdr_fuss)

        # Fusszeilen-Text: Firmenname · Tel · Fax · E-Mail · Internet
        fuss_text = (
            "{{fusszeile_firmenname}}"
            "  \u00b7  Tel: {{fusszeile_telefon}}"
            "  \u00b7  Fax: {{fusszeile_telefax}}"
            "  \u00b7  {{fusszeile_email}}"
            "  \u00b7  {{fusszeile_internet}}"
        )
        run_fuss = p_fuss.add_run(fuss_text)
        run_fuss.font.name = "Arial"
        run_fuss.font.size = Pt(8)
        run_fuss.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p_fuss.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Als Bytes zurueckgeben
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
