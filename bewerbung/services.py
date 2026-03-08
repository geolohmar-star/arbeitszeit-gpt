"""Services fuer die Bewerbungs-App:
  - stelle_ein(): Bewerbung -> HRMitarbeiter + Personalstammdaten
  - lehne_ab(): DSGVO Hard-Delete
  - erstelle_zusage_docx(): python-docx Zusage-Brief
  - erstelle_absage_docx(): python-docx Absage-Brief

Wird aufgerufen wenn HR einen Bewerber einstellt.
Erstellt automatisch: Django-User, HRMitarbeiter, Personalstammdaten.
Uebertraegt Dokumente in die verschluesselte dokumente-App.
Loescht anschliessend alle Bewerbungs-Rohdaten (DSGVO-Hard-Delete).
"""
import logging

from django.contrib.auth.models import User
from django.db import transaction

logger = logging.getLogger(__name__)


def _generiere_username(vorname: str, nachname: str) -> str:
    """Erstellt einen eindeutigen Username aus Vor- und Nachname."""
    import re
    basis = f"{vorname.lower()}.{nachname.lower()}"
    basis = re.sub(r"[^a-z0-9.]", "", basis.replace(" ", "."))
    username = basis
    zaehler = 1
    while User.objects.filter(username=username).exists():
        username = f"{basis}.{zaehler}"
        zaehler += 1
    return username


@transaction.atomic
def stelle_ein(bewerbung, stelle=None, eintrittsdatum=None, erstellt_von=None):
    """Stellt einen Bewerber ein.

    1. Django-User anlegen
    2. HRMitarbeiter anlegen (aus Bewerbungsdaten)
    3. Personalstammdaten anlegen (aus Bewerbungsdaten)
    4. Dokumente verschluesselt in dokumente-App uebertragen
    5. Bewerbungs-Rohdaten loeschen (DSGVO)

    Returns:
        hr.HRMitarbeiter – der neu angelegte Mitarbeiter
    """
    from hr.models import HRMitarbeiter, Personalstammdaten
    from dokumente.models import SensiblesDokument
    from dokumente.services import entschluessel_dokument, verschluessel_dokument

    # 1. Django-User erstellen
    username = _generiere_username(bewerbung.vorname, bewerbung.nachname)
    user = User.objects.create_user(
        username=username,
        first_name=bewerbung.vorname,
        last_name=bewerbung.nachname,
        email=bewerbung.email_privat or "",
    )
    logger.info("Einstellung: User '%s' erstellt.", username)

    # 2. HRMitarbeiter erstellen
    hr_ma = HRMitarbeiter.objects.create(
        user=user,
        vorname=bewerbung.vorname,
        nachname=bewerbung.nachname,
        stelle=stelle or bewerbung.angestrebte_stelle,
        eintrittsdatum=eintrittsdatum or bewerbung.geplantes_eintrittsdatum,
        email=bewerbung.email_privat or "",
    )
    logger.info("Einstellung: HRMitarbeiter '%s' (PNr: %s) erstellt.", hr_ma.vollname, hr_ma.personalnummer)

    # 3. Personalstammdaten erstellen (alle sensiblen Felder aus Bewerbung)
    Personalstammdaten.objects.create(
        mitarbeiter=hr_ma,
        angelegt_von=erstellt_von,
        anrede=bewerbung.anrede,
        geburtsdatum=bewerbung.geburtsdatum,
        geburtsort=bewerbung.geburtsort,
        geburtsname=bewerbung.geburtsname,
        staatsangehoerigkeit=bewerbung.staatsangehoerigkeit,
        familienstand=bewerbung.familienstand,
        konfession=bewerbung.konfession,
        anzahl_kinder=bewerbung.anzahl_kinder,
        strasse=bewerbung.strasse,
        hausnummer=bewerbung.hausnummer,
        plz=bewerbung.plz,
        ort=bewerbung.ort,
        land=bewerbung.land,
        telefon_privat=bewerbung.telefon_privat,
        mobil_privat=bewerbung.mobil_privat,
        email_privat=bewerbung.email_privat,
        steuerklasse=bewerbung.steuerklasse,
        steuer_id=bewerbung.steuer_id,
        sozialversicherungsnummer=bewerbung.sozialversicherungsnummer,
        iban=bewerbung.iban,
        bic=bewerbung.bic,
        bank_name=bewerbung.bank_name,
        krankenkasse_name=bewerbung.krankenkasse_name,
        krankenversicherungsart=bewerbung.krankenversicherungsart,
        notfallkontakt_name=bewerbung.notfallkontakt_name,
        notfallkontakt_beziehung=bewerbung.notfallkontakt_beziehung,
        notfallkontakt_telefon=bewerbung.notfallkontakt_telefon,
        vertragsart=bewerbung.vertragsart,
        probezeit_bis=bewerbung.probezeit_bis,
    )
    logger.info("Einstellung: Personalstammdaten fuer '%s' angelegt.", hr_ma.vollname)

    # 4. Bewerbungsdokumente in dokumente-App uebertragen (neu verschluesseln)
    kat_map = {
        "ausweis": "ausweis",
        "zeugnis_schule": "zeugnis",
        "zeugnis_arbeit": "zeugnis",
        "abschluss": "abschluss",
        "fuehrerschein": "ausweis",
        "sonstige": "sonstige",
    }
    for bdok in bewerbung.dokumente.all():
        try:
            inhalt_roh = entschluessel_dokument(bdok.inhalt_verschluesselt)
            inhalt_neu = verschluessel_dokument(inhalt_roh)
            SensiblesDokument.objects.create(
                user=user,
                hochgeladen_von=erstellt_von,
                kategorie=kat_map.get(bdok.typ, "sonstige"),
                dateiname=bdok.dateiname,
                dateityp=bdok.dateityp,
                inhalt_verschluesselt=inhalt_neu,
                groesse_bytes=bdok.groesse_bytes,
                beschreibung=f"Aus Bewerbung uebernommen ({bdok.get_typ_display()})",
            )
        except Exception as exc:
            logger.error("Dokument-Uebertragung fehlgeschlagen ('%s'): %s", bdok.dateiname, exc)

    logger.info("Einstellung: %d Dokument(e) uebertragen.", bewerbung.dokumente.count())

    # 5. Bewerbungs-Rohdaten loeschen (DSGVO Hard-Delete)
    bewerbung_name = bewerbung.vollname
    bewerbung.delete()
    logger.info("Einstellung abgeschlossen: Bewerbungsdaten '%s' geloescht (DSGVO).", bewerbung_name)

    return hr_ma


def erstelle_zusage_docx(bewerbung) -> bytes:
    """Erstellt einen Zusage-Brief als DOCX und gibt ihn als bytes zurueck."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.utils import timezone

    doc = Document()

    # Seitenraender
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Absenderzeile
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Personalabteilung – Interne Personalverwaltung")
    run.font.size = Pt(9)
    run.font.color.rgb = None

    doc.add_paragraph()

    # Empfaenger
    anrede_map = {"herr": "Herrn", "frau": "Frau", "divers": "", "keine": ""}
    anrede = anrede_map.get(bewerbung.anrede, "")
    empf = doc.add_paragraph()
    empf.add_run(f"{anrede} {bewerbung.vollname}\n".strip()).bold = True
    empf.add_run(
        f"{bewerbung.strasse} {bewerbung.hausnummer}\n"
        f"{bewerbung.plz} {bewerbung.ort}"
    )

    doc.add_paragraph()

    # Datum
    datum_p = doc.add_paragraph()
    datum_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    datum_p.add_run(timezone.localdate().strftime("%d.%m.%Y"))

    doc.add_paragraph()

    # Betreff
    betreff = doc.add_paragraph()
    r = betreff.add_run("Zusage Ihrer Bewerbung")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph()

    # Anredezeile
    anrede_brief = f"Sehr geehrte{'r' if bewerbung.anrede == 'herr' else ''} {anrede} {bewerbung.nachname}," if anrede else f"Sehr geehrte/r {bewerbung.vollname},"
    doc.add_paragraph(anrede_brief)

    doc.add_paragraph()

    # Text
    doc.add_paragraph(
        "wir freuen uns, Ihnen mitteilen zu koennen, dass wir Ihnen eine Stelle "
        "in unserem Unternehmen anbieten moechten."
    )

    # Stellendetails
    if bewerbung.angestrebte_stelle or bewerbung.geplantes_eintrittsdatum:
        doc.add_paragraph()
        tab = doc.add_table(rows=0, cols=2)
        tab.style = "Table Grid"
        if bewerbung.angestrebte_stelle:
            row = tab.add_row().cells
            row[0].text = "Stelle:"
            row[1].text = str(bewerbung.angestrebte_stelle)
        if bewerbung.geplantes_eintrittsdatum:
            row = tab.add_row().cells
            row[0].text = "Eintrittsdatum:"
            row[1].text = bewerbung.geplantes_eintrittsdatum.strftime("%d.%m.%Y")
        if bewerbung.vertragsart:
            row = tab.add_row().cells
            row[0].text = "Vertragsart:"
            row[1].text = bewerbung.get_vertragsart_display()

    doc.add_paragraph()
    doc.add_paragraph(
        "Bitte nehmen Sie zu den weiteren Einstellungsformalitaeten Kontakt "
        "mit der Personalabteilung auf. Wir benoetigen von Ihnen noch folgende "
        "Originaldokumente: Personalausweis, Sozialversicherungsnachweis, "
        "Steuer-ID sowie Ihre Bankverbindung."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Wir freuen uns auf Ihre Mitarbeit und wuenschen Ihnen einen guten Start."
    )
    doc.add_paragraph()
    doc.add_paragraph("Mit freundlichen Gruessen")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph("Personalabteilung")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def erstelle_absage_docx(bewerbung) -> bytes:
    """Erstellt einen Absage-Brief als DOCX und gibt ihn als bytes zurueck."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.utils import timezone

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    p = doc.add_paragraph()
    run = p.add_run("Personalabteilung – Interne Personalverwaltung")
    run.font.size = Pt(9)

    doc.add_paragraph()

    anrede_map = {"herr": "Herrn", "frau": "Frau", "divers": "", "keine": ""}
    anrede = anrede_map.get(bewerbung.anrede, "")
    empf = doc.add_paragraph()
    empf.add_run(f"{anrede} {bewerbung.vollname}\n".strip()).bold = True
    empf.add_run(
        f"{bewerbung.strasse} {bewerbung.hausnummer}\n"
        f"{bewerbung.plz} {bewerbung.ort}"
    )

    doc.add_paragraph()
    datum_p = doc.add_paragraph()
    datum_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    datum_p.add_run(timezone.localdate().strftime("%d.%m.%Y"))

    doc.add_paragraph()
    betreff = doc.add_paragraph()
    r = betreff.add_run("Ihre Bewerbung bei uns")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph()
    anrede_brief = f"Sehr geehrte{'r' if bewerbung.anrede == 'herr' else ''} {anrede} {bewerbung.nachname}," if anrede else f"Sehr geehrte/r {bewerbung.vollname},"
    doc.add_paragraph(anrede_brief)
    doc.add_paragraph()
    doc.add_paragraph(
        "vielen Dank fuer Ihr Interesse an einer Taetigkeit in unserem Unternehmen "
        "und die Zeit, die Sie sich fuer den Bewerbungsprozess genommen haben."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Nach sorgfaeltiger Pruefung Ihrer Unterlagen muessen wir Ihnen leider mitteilen, "
        "dass wir Ihre Bewerbung nicht beruecksichtigen koennen. Diese Entscheidung "
        "faellt uns nicht leicht und spiegelt nicht Ihre persoenlichen Qualitaeten wider."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Gemaess der Datenschutz-Grundverordnung (DSGVO) werden alle Ihre "
        "Bewerbungsunterlagen unverzueglich und vollstaendig geloescht."
    )
    doc.add_paragraph()
    doc.add_paragraph("Wir wuenschen Ihnen fuer Ihre weitere berufliche Zukunft alles Gute.")
    doc.add_paragraph()
    doc.add_paragraph("Mit freundlichen Gruessen")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph("Personalabteilung")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@transaction.atomic
def lehne_ab(bewerbung, abgelehnt_von=None):
    """Lehnt eine Bewerbung ab und loescht ALLE Daten unwiederbringlich (DSGVO).

    Kein Soft-Delete, kein Archiv – vollstaendige Vernichtung.
    """
    name = bewerbung.vollname
    bewerbung.delete()  # CASCADE loescht auch BewerbungDokument
    logger.info(
        "Bewerbung abgelehnt und DSGVO-konform geloescht: '%s' (abgelehnt von: %s).",
        name,
        abgelehnt_von.username if abgelehnt_von else "unbekannt",
    )
