"""
Django Views für Arbeitszeitverwaltung
"""
from django.db.models.functions import Length
import io
from collections import defaultdict
from datetime import datetime, date

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.db.models import Sum, Case, When, IntegerField
from django.http import HttpResponse
import csv
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from django.template.loader import render_to_string
from django.utils import timezone
# WICHTIG: Import für den Template-Filter-Workaround
from django.template.defaulttags import register

from .models import (
    Mitarbeiter,
    MonatlicheArbeitszeitSoll,
    Arbeitszeitvereinbarung,
    Tagesarbeitszeit,
    ArbeitszeitHistorie,
    Zeiterfassung,
    Urlaubsanspruch,
    Wochenbericht,
)
from .forms import RegisterForm

#WorkCalendar
from django.db.models import Sum, Avg, Count
from decimal import Decimal
import calendar
from .forms import SollStundenBerechnungForm
from django.utils.html import format_html
#Jahresübersicht Sol Stunden
import csv
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from weasyprint import HTML
from docx import Document

def hhmm_to_minuten(hhmm_wert):
    """Konvertiert HHMM-Format (z.B. 830 = 8:30) in Minuten."""
    if hhmm_wert is None:
        return 0
    stunden = hhmm_wert // 100
    minuten = hhmm_wert % 100
    return stunden * 60 + minuten

#WorkCalendar
@login_required
def soll_stunden_jahresuebersicht(request):
    """Jahresübersicht mit Export"""
    heute = timezone.now().date()
    jahr = int(request.GET.get('jahr', heute.year))
    export = request.GET.get('export')
    
    # Nur Mitarbeiter mit Schichtplan-Kennung MA1-MA15
    mitarbeiter_liste = Mitarbeiter.objects.filter(
        aktiv=True,
        schichtplan_kennung__regex=r'^MA([1-9]|1[0-5])$'
    ).exclude(
        schichtplan_kennung=''
    ).order_by(
        Length('schichtplan_kennung'),  # MA1 vor MA10
        'schichtplan_kennung'
    )
    
    # Daten sammeln
    daten = []
    
    for ma in mitarbeiter_liste:
        soll_monate = MonatlicheArbeitszeitSoll.objects.filter(
            mitarbeiter=ma,
            jahr=jahr
        ).order_by('monat')
        
        # Dict mit Monat -> Soll-Stunden
        monate_dict = {s.monat: float(s.soll_stunden) for s in soll_monate}
        jahressumme = sum(monate_dict.values())
        
        daten.append({
            'mitarbeiter': ma,
            'monate': monate_dict,
            'jahressumme': jahressumme,
            'durchschnitt': jahressumme / len(monate_dict) if monate_dict else 0,
        })
    
    # Export?
    if export == 'excel':
        return _export_excel(daten, jahr)
    elif export == 'csv':
        return _export_csv(daten, jahr)
    
    # Normal rendern
    context = {
        'jahr': jahr,
        'daten': daten,
        'monatsnamen': [calendar.month_abbr[i] for i in range(1, 13)],
        'vorheriges_jahr': jahr - 1,
        'naechstes_jahr': jahr + 1,
    }
    
    return render(request, 'arbeitszeit/soll_stunden_jahresuebersicht.html', context)


def _export_excel(daten, jahr):
    """Excel Export"""
    wb = Workbook()
    ws = wb.active
    ws.title = f"Soll-Stunden {jahr}"
    
    # Header
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    headers = ['Mitarbeiter', 'Kennung'] + \
              [calendar.month_abbr[i] for i in range(1, 13)] + \
              ['Summe', 'Durchschnitt']
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(1, col, header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    
    # Daten
    for row, item in enumerate(daten, 2):
        ws.cell(row, 1, item['mitarbeiter'].vollname)
        ws.cell(row, 2, item['mitarbeiter'].schichtplan_kennung or '-')
        
        for monat in range(1, 13):
            wert = item['monate'].get(monat, 0)
            ws.cell(row, monat + 2, wert)
        
        ws.cell(row, 15, item['jahressumme'])
        ws.cell(row, 16, item['durchschnitt'])
    
    # Spaltenbreiten
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 10
    
    # Response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=Soll-Stunden_{jahr}.xlsx'
    wb.save(response)
    return response


def _export_csv(daten, jahr):
    """CSV Export"""
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename=Soll-Stunden_{jahr}.csv'
    response.write('\ufeff')  # BOM für Excel
    
    writer = csv.writer(response, delimiter=';')
    
    # Header
    headers = ['Mitarbeiter', 'Kennung'] + \
              [calendar.month_name[i] for i in range(1, 13)] + \
              ['Summe', 'Durchschnitt']
    writer.writerow(headers)
    
    # Daten
    for item in daten:
        row = [
            item['mitarbeiter'].vollname,
            item['mitarbeiter'].schichtplan_kennung or '-',
        ]
        
        for monat in range(1, 13):
            wert = item['monate'].get(monat, 0)
            row.append(str(wert).replace('.', ','))
        
        row.append(str(item['jahressumme']).replace('.', ','))
        row.append(str(item['durchschnitt']).replace('.', ','))
        
        writer.writerow(row)
    
    return response


## ⚡ 4. Automatische Berechnung bei Login

### arbeitszeit/signals.py (NEU ERSTELLEN)

from django.contrib.auth.signals import user_logged_in
from django.dispatch import receiver
from django.utils import timezone
from .models import MonatlicheArbeitszeitSoll


@receiver(user_logged_in)
def auto_berechne_soll_stunden(sender, request, user, **kwargs):
    """Berechnet Soll-Stunden automatisch beim Login"""
    if not hasattr(user, 'mitarbeiter'):
        return
    
    heute = timezone.now().date()
    
    # Prüfe ob schon berechnet
    existiert = MonatlicheArbeitszeitSoll.objects.filter(
        mitarbeiter=user.mitarbeiter,
        jahr=heute.year,
        monat=heute.month
    ).exists()
    
    if not existiert:
        try:
            MonatlicheArbeitszeitSoll.berechne_und_speichere(
                user.mitarbeiter,
                heute.year,
                heute.month
            )
        except Exception as e:
            print(f"Auto-Berechnung fehlgeschlagen: {e}")
            
@login_required
def soll_stunden_dashboard(request):
    """
    Dashboard: Übersicht aller Soll-Stunden des aktuellen Monats
    """
    heute = timezone.now().date()
    jahr = request.GET.get('jahr', heute.year)
    monat = request.GET.get('monat', heute.month)
    
    try:
        jahr = int(jahr)
        monat = int(monat)
    except ValueError:
        jahr = heute.year
        monat = heute.month
    
    # Hole alle Soll-Stunden für diesen Monat
    soll_stunden_liste = MonatlicheArbeitszeitSoll.objects.filter(
        jahr=jahr,
        monat=monat,
        mitarbeiter__aktiv=True
    ).select_related('mitarbeiter').order_by('mitarbeiter__nachname')
    
    # Wenn keine Daten vorhanden, automatisch berechnen
    if not soll_stunden_liste.exists():
        messages.info(
            request,
            f"Keine Daten für {calendar.month_name[monat]} {jahr}. "
            f"Berechne automatisch..."
        )
        MonatlicheArbeitszeitSoll.berechne_fuer_alle_mitarbeiter(jahr, monat)
        
        # Neu laden
        soll_stunden_liste = MonatlicheArbeitszeitSoll.objects.filter(
            jahr=jahr,
            monat=monat,
            mitarbeiter__aktiv=True
        ).select_related('mitarbeiter').order_by('mitarbeiter__nachname')
    
    # Statistiken berechnen
    stats = soll_stunden_liste.aggregate(
        gesamt_soll=Sum('soll_stunden'),
        durchschnitt=Avg('soll_stunden'),
        anzahl=Count('id')
    )
    
    # Gesamtarbeitstage
    if soll_stunden_liste.exists():
        beispiel = soll_stunden_liste.first()
        arbeitstage_info = {
            'gesamt': beispiel.arbeitstage_gesamt,
            'feiertage': beispiel.feiertage_anzahl,
            'effektiv': beispiel.arbeitstage_effektiv,
            'feiertage_liste': beispiel.feiertage_liste
        }
    else:
        arbeitstage_info = None
    
    # Navigation: Vorheriger/Nächster Monat
    if monat == 1:
        vorheriger = {'jahr': jahr - 1, 'monat': 12}
    else:
        vorheriger = {'jahr': jahr, 'monat': monat - 1}
    
    if monat == 12:
        naechster = {'jahr': jahr + 1, 'monat': 1}
    else:
        naechster = {'jahr': jahr, 'monat': monat + 1}
    
    context = {
        'jahr': jahr,
        'monat': monat,
        'monat_name': calendar.month_name[monat],
        'soll_stunden_liste': soll_stunden_liste,
        'stats': stats,
        'arbeitstage_info': arbeitstage_info,
        'vorheriger': vorheriger,
        'naechster': naechster,
        'ist_aktueller_monat': (jahr == heute.year and monat == heute.month),
    }
    
    return render(request, 'arbeitszeit/soll_stunden_dashboard.html', context)


@login_required
def soll_stunden_berechnen(request):
    """
    View: Manuelle Berechnung der Soll-Stunden.
    Zeigt sanfte Warnungen für Mitarbeiter ohne Vereinbarung.
    """
    if request.method == 'POST':
        form = SollStundenBerechnungForm(request.POST)
        
        if form.is_valid():
            jahr = form.cleaned_data['jahr']
            monat = int(form.cleaned_data['monat'])
            mitarbeiter_qs = form.cleaned_data.get('mitarbeiter')
            
            if not mitarbeiter_qs:
                # Nur Mitarbeiter mit Schichtplan-Kennung MA1-MA15
                mitarbeiter_qs = Mitarbeiter.objects.filter(
                    aktiv=True,
                    schichtplan_kennung__regex=r'^MA([1-9]|1[0-5])$'  # MA1 bis MA15
                 
                ).exclude(
                    schichtplan_kennung=''  # Ausschließen ohne Kennung
                )
            
            # Zähler
            erfolge = 0
            fehler = []
            
            # Berechne für jeden Mitarbeiter
            for ma in mitarbeiter_qs:
                try:
                    # Versuche Berechnung
                    MonatlicheArbeitszeitSoll.berechne_und_speichere(ma, jahr, monat)
                    erfolge += 1
                    
                except ValueError as e:
                    # Keine Vereinbarung oder keine Wochenstunden
                    fehler.append({
                        'mitarbeiter': ma,
                        'fehler': str(e),
                        'typ': 'keine_vereinbarung'
                    })
                    
                except Exception as e:
                    # Anderer unerwarteter Fehler
                    fehler.append({
                        'mitarbeiter': ma,
                        'fehler': f"Unerwarteter Fehler: {str(e)}",
                        'typ': 'anderer_fehler'
                    })
            
            # Meldungen zusammenstellen
            monat_name = calendar.month_name[monat]
            
            # Erfolgsmeldung
            if erfolge > 0:
                messages.success(
                    request,
                    f"✅ Soll-Stunden für {erfolge} Mitarbeiter erfolgreich berechnet "
                    f"({monat_name} {jahr})"
                )
            
            # Warnungen für Fehler
            if fehler:
                # Gruppiere nach Typ
                ohne_vereinbarung = [f for f in fehler if f['typ'] == 'keine_vereinbarung']
                andere_fehler = [f for f in fehler if f['typ'] == 'anderer_fehler']
                
                if ohne_vereinbarung:
                    fehler_liste = "<ul>" + "".join([
                        f"<li><strong>{f['mitarbeiter'].nachname}, {f['mitarbeiter'].vorname}</strong> "
                        f"(ID: {f['mitarbeiter'].pk})</li>"
                        for f in ohne_vereinbarung
                    ]) + "</ul>"
                    
                    messages.warning(
                        request,
                        format_html(
                            "⚠️ <strong>{} Mitarbeiter ohne Arbeitszeitvereinbarung übersprungen:</strong><br>{}"
                            "<br><small>Bitte erstelle Vereinbarungen mit Status 'Aktiv' oder 'Genehmigt' "
                            "im <a href='/admin/arbeitszeit/arbeitszeitvereinbarung/add/' target='_blank'>Admin</a>.</small>",
                            len(ohne_vereinbarung),
                            fehler_liste
                        )
                    )
                
                if andere_fehler:
                    fehler_liste = "<ul>" + "".join([
                        f"<li><strong>{f['mitarbeiter'].nachname}, {f['mitarbeiter'].vorname}:</strong> "
                        f"{f['fehler']}</li>"
                        for f in andere_fehler
                    ]) + "</ul>"
                    
                    messages.error(
                        request,
                        format_html(
                            "❌ <strong>{} Mitarbeiter mit Fehlern:</strong><br>{}",
                            len(andere_fehler),
                            fehler_liste
                        )
                    )
            
            # Keine Erfolge und keine Fehler?
            if erfolge == 0 and not fehler:
                messages.info(
                    request,
                    "ℹ️ Keine Mitarbeiter für die Berechnung ausgewählt."
                )
            
            return redirect('arbeitszeit:soll_stunden_dashboard')
    
    else:
        form = SollStundenBerechnungForm()
    
    context = {
        'form': form,
    }
    
    return render(request, 'arbeitszeit/soll_stunden_berechnen.html', context)


@login_required
def mitarbeiter_soll_uebersicht(request, pk):
    """
    Detailansicht: Soll-Stunden eines Mitarbeiters über mehrere Monate
    """
    
    # Prüfe Berechtigung
    # Berechtigung: Staff ODER Schichtplaner
    if not (request.user.is_staff or 
            (hasattr(request.user, 'mitarbeiter') and 
             request.user.mitarbeiter.rolle == 'schichtplaner')):
        messages.error(request, "Keine Berechtigung.")
        return redirect('arbeitszeit:dashboard')
    
    mitarbeiter = get_object_or_404(Mitarbeiter, pk=pk)
    
    # Schichtplaner: Nur eigene Abteilung
    if (hasattr(request.user, 'mitarbeiter') and 
        request.user.mitarbeiter.rolle == 'schichtplaner' and
        mitarbeiter.abteilung != request.user.mitarbeiter.abteilung):
        messages.error(request, "Sie können nur Mitarbeiter Ihrer Abteilung einsehen.")
        return redirect('arbeitszeit:soll_stunden_dashboard')
    
    heute = timezone.now().date()
    
    jahr = int(request.GET.get('jahr', timezone.now().year))
    
    # Hole Soll-Stunden für das ganze Jahr
    soll_stunden_jahr = MonatlicheArbeitszeitSoll.objects.filter(
        mitarbeiter=mitarbeiter,
        jahr=jahr
    ).order_by('monat')
    
    # Wenn keine Daten vorhanden, für aktuellen Monat berechnen
    if not soll_stunden_jahr.exists():
        MonatlicheArbeitszeitSoll.berechne_und_speichere(
            mitarbeiter, 
            heute.year, 
            heute.month
        )
        soll_stunden_jahr = MonatlicheArbeitszeitSoll.objects.filter(
            mitarbeiter=mitarbeiter,
            jahr=jahr
        ).order_by('monat')
    
    # Statistiken
    stats = soll_stunden_jahr.aggregate(
        gesamt_jahr=Sum('soll_stunden'),
        durchschnitt=Avg('soll_stunden'),
        monate_berechnet=Count('id')
    )
    
    # Aktuelle Vereinbarung
    vereinbarung = mitarbeiter.get_aktuelle_vereinbarung()
    
    context = {
        'mitarbeiter': mitarbeiter,
        'jahr': jahr,
        'soll_stunden_jahr': soll_stunden_jahr,
        'stats': stats,
        'vereinbarung': vereinbarung,
        'vorheriges_jahr': jahr - 1,
        'naechstes_jahr': jahr + 1,
    }
    
    return render(request, 'arbeitszeit/mitarbeiter_soll_uebersicht.html', context)

#####Soll Stunden 
@login_required

def mitarbeiter_ohne_vereinbarung(request):
    """
    Zeigt alle Mitarbeiter ohne gültige Arbeitszeitvereinbarung.
    """
    heute = timezone.now().date()
    
    alle_mitarbeiter = Mitarbeiter.objects.filter(aktiv=True)
    ohne_vereinbarung = []
    
    for ma in alle_mitarbeiter:
        wochenstunden = ma.get_wochenstunden(heute)
        if wochenstunden is None:
            ohne_vereinbarung.append(ma)
    
    context = {
        'mitarbeiter_ohne_vereinbarung': ohne_vereinbarung,
        'anzahl': len(ohne_vereinbarung),
    }
    
    return render(request, 'arbeitszeit/mitarbeiter_ohne_vereinbarung.html', context)



# --- CUSTOM TEMPLATE FILTER (WORKAROUND) ---
@register.filter(name='mod')
def mod(value, arg):
    try:
        return int(value) % int(arg)
    except (ValueError, TypeError):
        return 0

# --- HELPER FUNKTIONEN ---

def _get_wochentag_sortierung():
    return Case(
        When(wochentag='montag', then=1),
        When(wochentag='dienstag', then=2),
        When(wochentag='mittwoch', then=3),
        When(wochentag='donnerstag', then=4),
        When(wochentag='freitag', then=5),
        When(wochentag='samstag', then=6),
        When(wochentag='sonntag', then=7),
        default=99,
        output_field=IntegerField(),
    )

def zeitwert_to_str(minuten):
    """Minuten in HH:MM umwandeln."""
    if minuten is None:
        return "0:00"
    stunden = minuten // 60
    rest_minuten = minuten % 60
    return f"{stunden}:{rest_minuten:02d}"

def get_zeitoptionen():
    zeitoptionen = [{'value': 0, 'label': 'Frei'}]
    for hour in range(2, 13):
        for minute in range(0, 60, 15):
            if hour == 12 and minute > 15:
                break
            value = hour * 60 + minute
            label = f"{hour:02d}:{minute:02d}"
            zeitoptionen.append({'value': value, 'label': label})
    return zeitoptionen

# --- VIEWS ---


def vereinbarung_detail(request, pk):
    vereinbarung = get_object_or_404(Arbeitszeitvereinbarung, pk=pk)
    arbeitszeiten = vereinbarung.tagesarbeitszeiten.annotate(
        sort_order=_get_wochentag_sortierung()
    ).order_by('woche', 'sort_order')
    
    temp_grouped = defaultdict(list)
    for ta in arbeitszeiten:
        # DEBUG: Zeige rohe Werte
        print(f"DEBUG: {ta.wochentag} - zeitwert RAW={ta.zeitwert}, in Minuten={hhmm_to_minuten(ta.zeitwert)}")
        ta.display_zeit = zeitwert_to_str(hhmm_to_minuten(ta.zeitwert))
        temp_grouped[ta.woche].append(ta)
    
    wochen_daten = {}
    for woche, tage in temp_grouped.items():
        gesamt_minuten = sum(hhmm_to_minuten(t.zeitwert) for t in tage if t.zeitwert)
        # DEBUG: Zeige Berechnung
        print(f"DEBUG Woche {woche}: gesamt_minuten={gesamt_minuten}, als String={zeitwert_to_str(gesamt_minuten)}")
        wochen_daten[woche] = {
            'tage': tage,
            'summe': zeitwert_to_str(gesamt_minuten)
        }
    
    historie = ArbeitszeitHistorie.objects.filter(
        vereinbarung=vereinbarung
    ).order_by('aenderung_am')

    # Vorherige/naechste Version ermitteln
    vorherige_version = Arbeitszeitvereinbarung.objects.filter(
        mitarbeiter=vereinbarung.mitarbeiter,
        versionsnummer__lt=vereinbarung.versionsnummer,
    ).order_by("-versionsnummer").first()
    naechste_version = Arbeitszeitvereinbarung.objects.filter(
        mitarbeiter=vereinbarung.mitarbeiter,
        versionsnummer__gt=vereinbarung.versionsnummer,
    ).order_by("versionsnummer").first()

    context = {
        'vereinbarung': vereinbarung,
        'wochen_daten': wochen_daten,
        'historie': historie,
        'vorherige_version': vorherige_version,
        'naechste_version': naechste_version,
    }
    return render(request, 'arbeitszeit/vereinbarung_detail.html', context)


@login_required
def dashboard(request):
    user = request.user
    
    # Sachbearbeiter und Schichtplaner sehen ihr persoenliches Dashboard
    # mit Link zum Admin-Bereich
    
    mitarbeiter, created = Mitarbeiter.objects.get_or_create(
        user=user,
        defaults={
            'personalnummer': f'MA{user.id:04d}',
            'vorname': user.first_name or 'Vorname',
            'nachname': user.last_name or 'Nachname',
            'abteilung': 'Allgemein',
            'standort': 'siegburg',
            'eintrittsdatum': date.today(),
            'aktiv': True
        }
    )
    if created:
        messages.info(request, "Ihr Mitarbeiter-Profil wurde erstellt.")
    
    aktuelle_vereinbarung = mitarbeiter.get_aktuelle_vereinbarung()
    letzte_erfassungen = mitarbeiter.zeiterfassungen.all().order_by('-datum')[:10]
    
    aktuelles_jahr = timezone.now().year
    urlaubsanspruch = mitarbeiter.urlaubsansprueche.filter(jahr=aktuelles_jahr).first()
    
    alle_vereinbarungen = mitarbeiter.arbeitszeitvereinbarungen.all().order_by('-created_at')[:5]
    
    is_kongos = (mitarbeiter.abteilung or '').strip().lower() == 'kongos'
    is_admin = request.user.is_staff or mitarbeiter.rolle == 'sachbearbeiter'
    context = {
        'mitarbeiter': mitarbeiter,
        'aktuelle_vereinbarung': aktuelle_vereinbarung,
        'letzte_erfassungen': letzte_erfassungen,
        'urlaubsanspruch': urlaubsanspruch,
        'alle_vereinbarungen': alle_vereinbarungen,
        'user': user,
        'is_kongos': is_kongos,
        'is_admin': is_admin,
    }
    return render(request, 'arbeitszeit/dashboard.html', context)


@login_required
def vereinbarung_erstellen(request):
    """Neue Arbeitszeitvereinbarung erstellen"""
    try:
        mitarbeiter = request.user.mitarbeiter
    except Mitarbeiter.DoesNotExist:
        messages.error(request, "Sie sind keinem Mitarbeiter zugeordnet.")
        return redirect('admin:index')
    
    # Tage-Liste für Template
    tage_list = ['montag', 'dienstag', 'mittwoch', 'donnerstag', 'freitag']  # ← NEU!
    
    if request.method == 'POST':
        antragsart = request.POST.get('antragsart')

        # Kettenmodell: Keine Ueberschneidungspruefung noetig.
        # Versionsnummer automatisch berechnen (hoechste + 1)
        letzte_version = mitarbeiter.arbeitszeitvereinbarungen.order_by(
            "-versionsnummer"
        ).values_list("versionsnummer", flat=True).first()
        naechste_version = (letzte_version or 0) + 1

        # Basisvereinbarung erstellen
        vereinbarung = Arbeitszeitvereinbarung(
            mitarbeiter=mitarbeiter,
            antragsart=antragsart,
            gueltig_ab=request.POST.get('gueltig_ab'),
            telearbeit=request.POST.get('telearbeit') == 'on',
            status='beantragt',
            versionsnummer=naechste_version,
        )
        
        # Bei Beendigung keine Arbeitszeit
        if antragsart == 'beendigung':
            vereinbarung.beendigung_beantragt = True
            vereinbarung.beendigung_datum = vereinbarung.gueltig_ab
        else:
            arbeitszeit_typ = request.POST.get('arbeitszeit_typ')
            vereinbarung.arbeitszeit_typ = arbeitszeit_typ
            
            if arbeitszeit_typ == 'regelmaessig':
                vereinbarung.wochenstunden = request.POST.get('wochenstunden')
        
        vereinbarung.save()
        
        # Tagesarbeitszeiten speichern (individuell)
        if antragsart != 'beendigung' and request.POST.get('arbeitszeit_typ') == 'individuell':
            gesamt_minuten = 0
            
            # Suche nach Wochen (neuantrag_montag_1, neuantrag_montag_2, etc.)
            week = 1
            while True:
                found_in_week = False
                
                for tag in tage_list:
                    # Suche nach neuantrag_montag_1 (MIT _1, _2, etc.)
                    zeitwert_str = request.POST.get(f'neuantrag_{tag}_{week}')
                    
                    if zeitwert_str:
                        found_in_week = True
                        try:
                            # Erwarte HH:MM Format vom type="time" Input
                            if ':' in zeitwert_str:
                                stunden, minuten = map(int, zeitwert_str.split(':'))
                            else:
                                # Fallback: HMM Format
                                wert = int(zeitwert_str)
                                stunden = wert // 100
                                minuten = wert % 100
                            
                            wert_minuten = stunden * 60 + minuten
                            
                            # Speichere als HMM für Kompatibilität
                            wert_hmm = stunden * 100 + minuten
                            
                            Tagesarbeitszeit.objects.create(
                                vereinbarung=vereinbarung,
                                wochentag=tag,
                                zeitwert=wert_hmm
                            )
                            
                            gesamt_minuten += wert_minuten
                            
                        except (ValueError, AttributeError) as e:
                            print(f"Fehler bei {tag} Woche {week}: {e}")
                            continue
                
                # Keine Daten mehr in dieser Woche? → Abbruch
                if not found_in_week:
                    break
                    
                week += 1
            
            # Wochenstunden berechnen
            if gesamt_minuten > 0:
                vereinbarung.wochenstunden = round(gesamt_minuten / 60, 2)
                vereinbarung.save()
        
        # Gültigkeit
        if antragsart != 'beendigung':
            gueltig_bis = request.POST.get('gueltig_bis')
            if gueltig_bis:
                vereinbarung.gueltig_bis = gueltig_bis
        
        vereinbarung.save()
        
        # Historie
        ArbeitszeitHistorie.objects.create(
            vereinbarung=vereinbarung,
            aenderung_durch=request.user,
            alter_status='entwurf',
            neuer_status='beantragt',
            bemerkung=f'{vereinbarung.get_antragsart_display()} erstellt und beantragt'
        )
        
        messages.success(request, 'Arbeitszeitvereinbarung wurde erfolgreich beantragt.')
        return redirect('arbeitszeit:dashboard')
    
    context = {
        'mitarbeiter': mitarbeiter,
        'tage_list': tage_list,  # ← NEU!
    }
    
    return render(request, 'arbeitszeit/vereinbarung_form.html', context)


@login_required
def vereinbarung_liste(request):
    try:
        mitarbeiter = request.user.mitarbeiter
    except Mitarbeiter.DoesNotExist:
        return redirect('arbeitszeit:dashboard')

    vereinbarungen = mitarbeiter.arbeitszeitvereinbarungen.all().order_by(
        '-gueltig_ab', '-versionsnummer'
    )

    status_filter = request.GET.get('status')
    if status_filter:
        vereinbarungen = vereinbarungen.filter(status=status_filter)

    # Aktuelle Version markieren (die zum heutigen Datum gilt)
    aktuelle = mitarbeiter.get_aktuelle_vereinbarung()
    aktuelle_pk = aktuelle.pk if aktuelle else None

    context = {
        'mitarbeiter': mitarbeiter,
        'vereinbarungen': vereinbarungen,
        'status_filter': status_filter,
        'aktuelle_pk': aktuelle_pk,
    }
    return render(request, 'arbeitszeit/vereinbarung_liste.html', context)


def _soll_minuten_aus_vereinbarung(mitarbeiter, datum):
    """Ermittelt Soll-Minuten aus der gueltigen Vereinbarung.

    - An Feiertagen (standortabhaengig) wird 0 zurueckgegeben.
    - Bei individueller Vereinbarung: tatsaechliche Tagesarbeitszeit.
    - Bei Wechselwochen: Durchschnitt ueber alle Wochen.
    - Bei regelmaessiger Vereinbarung: Wochenstunden / 5.
    """
    from .models import get_feiertagskalender, Tagesarbeitszeit

    # Feiertags-Check: kein Soll an Feiertagen
    cal = get_feiertagskalender(mitarbeiter.standort)
    if cal.is_holiday(datum):
        return 0

    vereinbarung = mitarbeiter.get_aktuelle_vereinbarung(datum)
    if not vereinbarung:
        return None

    # Wochentag-Mapping: Python weekday() -> Tagesarbeitszeit
    WOCHENTAG_MAP = {
        0: "montag",
        1: "dienstag",
        2: "mittwoch",
        3: "donnerstag",
        4: "freitag",
        5: "samstag",
        6: "sonntag",
    }
    wochentag_name = WOCHENTAG_MAP[datum.weekday()]

    if vereinbarung.arbeitszeit_typ == "individuell":
        # Tagesarbeitszeiten fuer diesen Wochentag holen
        tage = Tagesarbeitszeit.objects.filter(
            vereinbarung=vereinbarung,
            wochentag=wochentag_name,
        )
        if tage.exists():
            # Durchschnitt ueber alle Wochen (z.B. Woche 1 + 2)
            gesamt = sum(t.zeit_in_minuten for t in tage)
            return int(round(gesamt / tage.count()))
        # Kein Eintrag fuer diesen Tag -> 0 Soll
        return 0

    # Regelmaessig: Wochenstunden / 5
    if vereinbarung.wochenstunden:
        tages_soll = float(vereinbarung.wochenstunden) / 5
        return int(round(tages_soll * 60))
    return None


def _erstelle_urlaub_eintraege(mitarbeiter, datum_von, datum_bis,
                               soll_minuten, bemerkung):
    """Erstellt Zeiterfassungs-Eintraege fuer Urlaubszeitraum.

    Wochenenden und Feiertage (standortabhaengig) werden uebersprungen.
    """
    from datetime import timedelta
    from .models import get_feiertagskalender

    cal = get_feiertagskalender(mitarbeiter.standort)

    aktuell = datum_von
    anzahl = 0
    while aktuell <= datum_bis:
        # Nur Werktage (Mo-Fr) und keine Feiertage
        if aktuell.weekday() < 5 and not cal.is_holiday(aktuell):
            Zeiterfassung.objects.update_or_create(
                mitarbeiter=mitarbeiter,
                datum=aktuell,
                defaults={
                    "art": "urlaub",
                    "arbeitsbeginn": None,
                    "arbeitsende": None,
                    "pause_minuten": 0,
                    "arbeitszeit_minuten": None,
                    "soll_minuten": soll_minuten,
                    "urlaub_bis": datum_bis if aktuell == datum_von
                    else None,
                    "bemerkung": bemerkung,
                },
            )
            anzahl += 1
        aktuell += timedelta(days=1)
    return anzahl


@login_required
def zeiterfassung_erstellen(request):
    try:
        mitarbeiter = request.user.mitarbeiter
    except Mitarbeiter.DoesNotExist:
        return redirect("arbeitszeit:dashboard")

    datum = (
        request.POST.get("datum")
        or request.GET.get("datum")
        or str(timezone.now().date())
    )

    if request.method == "POST":
        art = request.POST.get("art", "homeoffice")

        # Datum parsen
        try:
            datum_obj = date.fromisoformat(datum)
        except (ValueError, TypeError):
            datum_obj = timezone.now().date()

        # Soll-Minuten aus Vereinbarung
        soll_minuten = _soll_minuten_aus_vereinbarung(
            mitarbeiter, datum_obj
        )

        bemerkung = request.POST.get("bemerkung", "")

        if art == "urlaub":
            # Urlaub: optional Datumsbereich
            urlaub_bis_str = request.POST.get("urlaub_bis", "")
            if urlaub_bis_str:
                try:
                    urlaub_bis_obj = date.fromisoformat(urlaub_bis_str)
                except (ValueError, TypeError):
                    urlaub_bis_obj = datum_obj
            else:
                urlaub_bis_obj = datum_obj

            if urlaub_bis_obj < datum_obj:
                urlaub_bis_obj = datum_obj

            anzahl = _erstelle_urlaub_eintraege(
                mitarbeiter, datum_obj, urlaub_bis_obj,
                soll_minuten, bemerkung,
            )
            messages.success(
                request,
                f"Urlaub fuer {anzahl} Tag(e) eingetragen.",
            )
            # Zurueck zur KW des erfassten Datums
            iso = datum_obj.isocalendar()
            return redirect(
                f"/zeiterfassung/?ansicht=woche"
                f"&kw={iso[1]}&jahr={iso[0]}"
            )

        elif art in ("krank", "z_ag"):
            # Krank / Z-AG: keine Zeitfelder
            defaults = {
                "art": art,
                "arbeitsbeginn": None,
                "arbeitsende": None,
                "pause_minuten": 0,
                "soll_minuten": soll_minuten,
                "bemerkung": bemerkung,
            }
            if art == "z_ag":
                # Z-AG: keine Arbeitszeit, Differenz ergibt -Soll
                defaults["arbeitszeit_minuten"] = 0
            elif art == "krank":
                # Krank: Soll wird gutgeschrieben, kein Abzug
                defaults["arbeitszeit_minuten"] = (
                    soll_minuten if soll_minuten else 0
                )

            Zeiterfassung.objects.update_or_create(
                mitarbeiter=mitarbeiter,
                datum=datum_obj,
                defaults=defaults,
            )
            messages.success(
                request, "Zeiterfassung wurde gespeichert."
            )
            iso = datum_obj.isocalendar()
            return redirect(
                f"/zeiterfassung/?ansicht=woche"
                f"&kw={iso[1]}&jahr={iso[0]}"
            )

        else:
            # HomeOffice / Telearbeit / Hybrid: mit Zeitfeldern
            from datetime import time as dt_time

            beginn_str = request.POST.get("arbeitsbeginn", "")
            ende_str = request.POST.get("arbeitsende", "")
            arbeitsbeginn = None
            arbeitsende = None
            if beginn_str:
                h, m = beginn_str.split(":")
                arbeitsbeginn = dt_time(int(h), int(m))
            if ende_str:
                h, m = ende_str.split(":")
                arbeitsende = dt_time(int(h), int(m))

            # Manuelle Pause (optional)
            manuelle_pause_str = request.POST.get(
                "manuelle_pause", ""
            )
            manuelle_pause = None
            if manuelle_pause_str.strip():
                try:
                    manuelle_pause = int(manuelle_pause_str)
                except (ValueError, TypeError):
                    manuelle_pause = None

            Zeiterfassung.objects.update_or_create(
                mitarbeiter=mitarbeiter,
                datum=datum_obj,
                defaults={
                    "art": art,
                    "arbeitsbeginn": arbeitsbeginn,
                    "arbeitsende": arbeitsende,
                    "manuelle_pause": manuelle_pause,
                    "soll_minuten": soll_minuten,
                    "bemerkung": bemerkung,
                },
            )
            messages.success(
                request, "Zeiterfassung wurde gespeichert."
            )
            iso = datum_obj.isocalendar()
            return redirect(
                f"/zeiterfassung/?ansicht=woche"
                f"&kw={iso[1]}&jahr={iso[0]}"
            )

    # GET: Bestehende Erfassung laden
    if isinstance(datum, str):
        try:
            datum_obj = date.fromisoformat(datum)
        except (ValueError, TypeError):
            datum_obj = timezone.now().date()
    else:
        datum_obj = datum

    erfassung = Zeiterfassung.objects.filter(
        mitarbeiter=mitarbeiter, datum=datum_obj
    ).first()

    # Soll-Info fuer Template
    soll_minuten = _soll_minuten_aus_vereinbarung(
        mitarbeiter, datum_obj
    )
    vereinbarung = mitarbeiter.get_aktuelle_vereinbarung(datum_obj)

    # Formatierte Werte fuer Template
    brutto_formatiert = None
    differenz_formatiert = None
    if erfassung and erfassung.brutto_minuten:
        b = erfassung.brutto_minuten
        brutto_formatiert = f"{b // 60}:{b % 60:02d}h ({b} min)"
    if erfassung and erfassung.differenz_minuten is not None:
        diff = erfassung.differenz_minuten
        vorzeichen = "+" if diff >= 0 else ""
        abs_diff = abs(diff)
        differenz_formatiert = (
            f"{vorzeichen}{abs_diff // 60}:{abs_diff % 60:02d}h"
            f" ({vorzeichen}{diff} min)"
        )

    context = {
        "mitarbeiter": mitarbeiter,
        "datum": datum_obj,
        "erfassung": erfassung,
        "soll_minuten": soll_minuten,
        "soll_formatiert": (
            f"{soll_minuten // 60}:{soll_minuten % 60:02d}h"
            if soll_minuten
            else None
        ),
        "vereinbarung": vereinbarung,
        "brutto_formatiert": brutto_formatiert,
        "differenz_formatiert": differenz_formatiert,
    }
    return render(
        request, "arbeitszeit/zeiterfassung_form.html", context
    )


def _minuten_formatiert(minuten):
    """Hilfsfunktion: Minuten als +/-H:MMh formatieren."""
    abs_m = abs(minuten)
    vz = "+" if minuten >= 0 else "-"
    return f"{vz}{abs_m // 60}:{abs_m % 60:02d}h"


@login_required
def zeiterfassung_uebersicht(request):
    try:
        mitarbeiter = request.user.mitarbeiter
    except Mitarbeiter.DoesNotExist:
        return redirect("arbeitszeit:dashboard")

    from datetime import timedelta

    # Ansicht: "woche" oder "monat"
    ansicht = request.GET.get("ansicht", "woche")

    heute = timezone.now().date()
    jahr = int(request.GET.get("jahr", heute.year))

    if ansicht == "woche":
        # KW ermitteln (ISO-Kalenderwoche)
        kw = int(request.GET.get("kw", heute.isocalendar()[1]))

        # Montag der KW berechnen
        # ISO: Woche 1 enthaelt den 4. Januar
        jan4 = date(jahr, 1, 4)
        # Montag der KW 1
        montag_kw1 = jan4 - timedelta(days=jan4.weekday())
        montag = montag_kw1 + timedelta(weeks=kw - 1)
        sonntag = montag + timedelta(days=6)

        # Vorherige / naechste KW
        prev_montag = montag - timedelta(weeks=1)
        next_montag = montag + timedelta(weeks=1)
        prev_kw = prev_montag.isocalendar()[1]
        prev_jahr = prev_montag.isocalendar()[0]
        next_kw = next_montag.isocalendar()[1]
        next_jahr = next_montag.isocalendar()[0]

        # Erfassungen der Woche laden
        erfassungen_qs = mitarbeiter.zeiterfassungen.filter(
            datum__gte=montag, datum__lte=sonntag,
        ).order_by("datum")

        # Dict fuer schnellen Zugriff nach Datum
        erfassungen_dict = {
            e.datum: e for e in erfassungen_qs
        }

        # Feiertagskalender fuer Standort
        from .models import get_feiertagskalender, feiertag_name_deutsch
        cal = get_feiertagskalender(mitarbeiter.standort)

        # 7 Tage Mo-So aufbauen (auch leere)
        WOCHENTAGE = [
            "Mo", "Di", "Mi", "Do", "Fr", "Sa", "So",
        ]
        wochen_tage = []
        for i in range(7):
            tag_datum = montag + timedelta(days=i)
            erfassung = erfassungen_dict.get(tag_datum)
            ist_feiertag = cal.is_holiday(tag_datum)
            feiertag_name = (
                feiertag_name_deutsch(cal, tag_datum)
                if ist_feiertag else ""
            )
            wochen_tage.append({
                "datum": tag_datum,
                "wochentag": WOCHENTAGE[i],
                "ist_heute": tag_datum == heute,
                "ist_wochenende": i >= 5,
                "ist_feiertag": ist_feiertag,
                "feiertag_name": feiertag_name,
                "erfassung": erfassung,
            })

        # Summen (Urlaub wird nicht mitgerechnet)
        erfassungen_ohne_urlaub = erfassungen_qs.exclude(
            art="urlaub"
        )
        gesamt_minuten = (
            erfassungen_ohne_urlaub.aggregate(
                total=Sum("arbeitszeit_minuten")
            )["total"]
            or 0
        )
        gesamt_differenz = sum(
            e.differenz_minuten for e in erfassungen_ohne_urlaub
            if e.differenz_minuten is not None
        )

        # Wochenbericht-Status abfragen
        wochenbericht = Wochenbericht.objects.filter(
            mitarbeiter=mitarbeiter, jahr=jahr, kw=kw,
        ).first()

        context = {
            "mitarbeiter": mitarbeiter,
            "ansicht": "woche",
            "jahr": jahr,
            "kw": kw,
            "montag": montag,
            "sonntag": sonntag,
            "wochen_tage": wochen_tage,
            "erfassungen": erfassungen_qs,
            "gesamt_arbeitszeit": (
                zeitwert_to_str(gesamt_minuten) + "h"
            ),
            "gesamt_differenz": gesamt_differenz,
            "gesamt_differenz_formatiert": (
                _minuten_formatiert(gesamt_differenz)
            ),
            "prev_kw": prev_kw,
            "prev_jahr": prev_jahr,
            "next_kw": next_kw,
            "next_jahr": next_jahr,
            "wochenbericht": wochenbericht,
        }

    else:
        # Monatsansicht (bestehend)
        monat = int(request.GET.get("monat", heute.month))

        # Feiertagskalender fuer Standort
        from .models import get_feiertagskalender, feiertag_name_deutsch
        cal = get_feiertagskalender(mitarbeiter.standort)

        erfassungen = mitarbeiter.zeiterfassungen.filter(
            datum__year=jahr, datum__month=monat,
        ).order_by("datum")

        # Feiertage des Monats als Dict {datum: name}
        import datetime as dt_mod
        _, letzter_tag = calendar.monthrange(jahr, monat)
        feiertage_monat = {}
        for tag_nr in range(1, letzter_tag + 1):
            d = dt_mod.date(jahr, monat, tag_nr)
            if cal.is_holiday(d):
                feiertage_monat[d] = feiertag_name_deutsch(cal, d)

        # Summen (Urlaub wird nicht mitgerechnet)
        erfassungen_ohne_urlaub = erfassungen.exclude(
            art="urlaub"
        )
        gesamt_minuten = (
            erfassungen_ohne_urlaub.aggregate(
                total=Sum("arbeitszeit_minuten")
            )["total"]
            or 0
        )
        gesamt_differenz = sum(
            e.differenz_minuten for e in erfassungen_ohne_urlaub
            if e.differenz_minuten is not None
        )

        context = {
            "mitarbeiter": mitarbeiter,
            "ansicht": "monat",
            "erfassungen": erfassungen,
            "jahr": jahr,
            "monat": monat,
            "feiertage_monat": feiertage_monat,
            "gesamt_arbeitszeit": (
                zeitwert_to_str(gesamt_minuten) + "h"
            ),
            "gesamt_differenz": gesamt_differenz,
            "gesamt_differenz_formatiert": (
                _minuten_formatiert(gesamt_differenz)
            ),
        }

    return render(
        request,
        "arbeitszeit/zeiterfassung_uebersicht.html",
        context,
    )


@login_required
def zeiterfassung_loeschen(request, pk):
    """Loescht eine einzelne Zeiterfassung des eingeloggten Nutzers."""
    try:
        mitarbeiter = request.user.mitarbeiter
    except Mitarbeiter.DoesNotExist:
        return redirect("arbeitszeit:dashboard")

    erfassung = get_object_or_404(
        Zeiterfassung, pk=pk, mitarbeiter=mitarbeiter
    )

    if request.method == "POST":
        kw = erfassung.datum.isocalendar()[1]
        jahr = erfassung.datum.isocalendar()[0]
        erfassung.delete()
        messages.success(request, "Zeiterfassung wurde geloescht.")
        return redirect(
            f"/zeiterfassung/?ansicht=woche&kw={kw}&jahr={jahr}"
        )

    return redirect("arbeitszeit:zeiterfassung_uebersicht")


@login_required
def wochenbericht_pdf(request):
    """Generiert einen PDF-Wochenbericht fuer die angegebene KW."""
    try:
        mitarbeiter = request.user.mitarbeiter
    except Mitarbeiter.DoesNotExist:
        return redirect("arbeitszeit:dashboard")

    from datetime import timedelta
    from weasyprint import HTML
    from .models import get_feiertagskalender, feiertag_name_deutsch

    heute = timezone.now().date()
    jahr = int(request.GET.get("jahr", heute.year))
    kw = int(request.GET.get("kw", heute.isocalendar()[1]))

    # Montag der KW berechnen (gleiche Logik wie Uebersicht)
    jan4 = date(jahr, 1, 4)
    montag_kw1 = jan4 - timedelta(days=jan4.weekday())
    montag = montag_kw1 + timedelta(weeks=kw - 1)
    sonntag = montag + timedelta(days=6)

    # Erfassungen laden
    erfassungen_qs = mitarbeiter.zeiterfassungen.filter(
        datum__gte=montag, datum__lte=sonntag,
    ).order_by("datum")
    erfassungen_dict = {e.datum: e for e in erfassungen_qs}

    # Feiertagskalender
    cal = get_feiertagskalender(mitarbeiter.standort)

    WOCHENTAGE = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]
    wochen_tage = []
    for i in range(7):
        tag_datum = montag + timedelta(days=i)
        erfassung = erfassungen_dict.get(tag_datum)
        ist_feiertag = cal.is_holiday(tag_datum)
        feiertag_name = (
            feiertag_name_deutsch(cal, tag_datum)
            if ist_feiertag else ""
        )
        wochen_tage.append({
            "datum": tag_datum,
            "wochentag": WOCHENTAGE[i],
            "ist_wochenende": i >= 5,
            "ist_feiertag": ist_feiertag,
            "feiertag_name": feiertag_name,
            "erfassung": erfassung,
        })

    # Summen (Urlaub wird nicht mitgerechnet)
    erfassungen_ohne_urlaub = erfassungen_qs.exclude(
        art="urlaub"
    )
    gesamt_minuten = (
        erfassungen_ohne_urlaub.aggregate(
            total=Sum("arbeitszeit_minuten")
        )["total"]
        or 0
    )
    gesamt_differenz = sum(
        e.differenz_minuten for e in erfassungen_ohne_urlaub
        if e.differenz_minuten is not None
    )

    druckdatum = timezone.now()

    context = {
        "mitarbeiter": mitarbeiter,
        "jahr": jahr,
        "kw": kw,
        "montag": montag,
        "sonntag": sonntag,
        "wochen_tage": wochen_tage,
        "gesamt_arbeitszeit": zeitwert_to_str(gesamt_minuten) + "h",
        "gesamt_differenz": gesamt_differenz,
        "gesamt_differenz_formatiert": _minuten_formatiert(
            gesamt_differenz
        ),
        "druckdatum": druckdatum,
    }

    html_string = render_to_string(
        "arbeitszeit/pdf_wochenbericht.html", context
    )
    html = HTML(
        string=html_string,
        base_url=request.build_absolute_uri("/"),
    )
    pdf = html.write_pdf()

    # Druckstatus tracken
    Wochenbericht.objects.update_or_create(
        mitarbeiter=mitarbeiter,
        jahr=jahr,
        kw=kw,
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    filename = (
        f"Wochenbericht_{mitarbeiter.nachname}"
        f"_KW{kw}_{jahr}.pdf"
    )
    response["Content-Disposition"] = (
        f'attachment; filename="{filename}"'
    )
    return response


# --- ADMIN VIEWS ---

@login_required
def admin_dashboard(request):
    is_sachbearbeiter = (
        hasattr(request.user, 'mitarbeiter')
        and request.user.mitarbeiter.rolle == 'sachbearbeiter'
    )
    if not (request.user.is_staff or is_sachbearbeiter):
        messages.error(request, "Sie haben keine Berechtigung fuer diese Seite.")
        return redirect('arbeitszeit:dashboard')
    
    heute = timezone.now()
    monatsanfang = heute.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    
    stats = {
        'offene_antraege': Arbeitszeitvereinbarung.objects.filter(status='beantragt').count(),
        'gesamt_mitarbeiter': Mitarbeiter.objects.filter(aktiv=True).count(),
        'genehmigte_vereinbarungen': Arbeitszeitvereinbarung.objects.filter(
            status='genehmigt', genehmigt_am__gte=monatsanfang
        ).count(),
        'aktive_vereinbarungen': Arbeitszeitvereinbarung.objects.filter(status='aktiv').count(),
        'mitarbeiter_siegburg': Mitarbeiter.objects.filter(aktiv=True, standort='siegburg').count(),
        'mitarbeiter_bonn': Mitarbeiter.objects.filter(aktiv=True, standort='bonn').count(),
    }
    
    letzte_aktivitaeten = Arbeitszeitvereinbarung.objects.filter(
        created_at__gte=heute.date()
    ).select_related('mitarbeiter').order_by('-created_at')[:5]
    
    context = {**stats, 'letzte_aktivitaeten': letzte_aktivitaeten}
    return render(request, 'arbeitszeit/admin_dashboard.html', context)


@login_required
def admin_vereinbarungen_genehmigen(request):
    is_sachbearbeiter = (
        hasattr(request.user, 'mitarbeiter')
        and request.user.mitarbeiter.rolle == 'sachbearbeiter'
    )
    if not (request.user.is_staff or is_sachbearbeiter):
        return redirect('arbeitszeit:dashboard')

    from .filters import ArbeitszeitvereinbarungFilter

    offene_antraege = Arbeitszeitvereinbarung.objects.filter(
        status='beantragt'
    ).select_related('mitarbeiter').order_by('-created_at')

    alle_qs = Arbeitszeitvereinbarung.objects.select_related(
        'mitarbeiter'
    ).order_by('-created_at')

    vereinbarung_filter = ArbeitszeitvereinbarungFilter(
        request.GET, queryset=alle_qs
    )

    context = {
        'offene_antraege': offene_antraege,
        'alle_vereinbarungen': vereinbarung_filter.qs,
        'filter': vereinbarung_filter,
    }
    return render(request, 'arbeitszeit/admin_vereinbarungen.html', context)


@login_required
def admin_vereinbarung_genehmigen(request, pk):
    """Einzelne Vereinbarung genehmigen (Detailansicht fuer Admin)"""
    is_sachbearbeiter = (
        hasattr(request.user, 'mitarbeiter')
        and request.user.mitarbeiter.rolle == 'sachbearbeiter'
    )
    if not (request.user.is_staff or is_sachbearbeiter):
        return redirect('arbeitszeit:dashboard')
    
    vereinbarung = get_object_or_404(Arbeitszeitvereinbarung, pk=pk)
    
    if request.method == 'POST':
        aktion = request.POST.get('aktion')
        bemerkung = request.POST.get('bemerkung', '')
        alter_status = vereinbarung.status
        
        if aktion == 'genehmigen':
            vereinbarung.status = 'aktiv'
            vereinbarung.genehmigt_von = request.user
            vereinbarung.genehmigt_am = timezone.now()
            messages.success(request, 'Vereinbarung wurde genehmigt und aktiviert.')
        elif aktion == 'ablehnen':
            vereinbarung.status = 'abgelehnt'
            messages.warning(request, 'Vereinbarung wurde abgelehnt.')
        
        vereinbarung.save()
        
        ArbeitszeitHistorie.objects.create(
            vereinbarung=vereinbarung,
            aenderung_durch=request.user,
            alter_status=alter_status,
            neuer_status=vereinbarung.status,
            bemerkung=bemerkung
        )
        return redirect('arbeitszeit:admin_vereinbarungen')
    
    # Sortierung mittels Helper
    tagesarbeitszeiten = vereinbarung.tagesarbeitszeiten.annotate(
        sort_order=_get_wochentag_sortierung()
    ).order_by('woche', 'sort_order')
    
   
    temp_grouped = defaultdict(list)
    for tag in tagesarbeitszeiten:
        # DEBUG: Zeige rohe Werte in der Konsole
        print(f"DEBUG: Wochentag={tag.wochentag}, zeitwert RAW={tag.zeitwert}")
        
        # Konvertierung
        minuten = hhmm_to_minuten(tag.zeitwert)
        print(f"  -> Nach Konvertierung: {minuten} Minuten = {zeitwert_to_str(minuten)}")
    
        tag.display_zeit = zeitwert_to_str(minuten)
        temp_grouped[tag.woche].append(tag)

    
    
    
    # Auch hier: Strukturierte Daten mit Summe
    wochen_daten = {}
    for woche, tage in temp_grouped.items():
        # DEBUG: Zeige Einzelwerte
        print(f"\nDEBUG Woche {woche}:")
        for t in tage:
            print(f"  - {t.wochentag}: zeitwert={t.zeitwert}")
        
        gesamt_minuten = sum(hhmm_to_minuten(t.zeitwert) for t in tage if t.zeitwert)
        print(f"  -> SUMME: {gesamt_minuten} Minuten = {zeitwert_to_str(gesamt_minuten)}\n")
        
        wochen_daten[woche] = {
            'tage': tage,
            'summe': zeitwert_to_str(gesamt_minuten)
        }

    # Vorherige/naechste Version ermitteln
    vorherige_version = Arbeitszeitvereinbarung.objects.filter(
        mitarbeiter=vereinbarung.mitarbeiter,
        versionsnummer__lt=vereinbarung.versionsnummer,
    ).order_by("-versionsnummer").first()
    naechste_version = Arbeitszeitvereinbarung.objects.filter(
        mitarbeiter=vereinbarung.mitarbeiter,
        versionsnummer__gt=vereinbarung.versionsnummer,
    ).order_by("versionsnummer").first()

    context = {
        'vereinbarung': vereinbarung,
        'wochen_daten': wochen_daten,
        'zeitoptionen': get_zeitoptionen(),
        'vorherige_version': vorherige_version,
        'naechste_version': naechste_version,
    }
    return render(request, 'arbeitszeit/admin_vereinbarung_detail.html', context)


@login_required
def mitarbeiter_uebersicht(request):
    is_schichtplaner = hasattr(request.user, 'mitarbeiter') and request.user.mitarbeiter.rolle == 'schichtplaner'
    
    if not (request.user.is_staff or is_schichtplaner):
        messages.error(request, "Keine Berechtigung.")
        return redirect('arbeitszeit:dashboard')
    
    mitarbeiter_liste = Mitarbeiter.objects.filter(aktiv=True)
    
    if not request.user.is_staff and is_schichtplaner:
        eigene_abteilung = request.user.mitarbeiter.abteilung
        mitarbeiter_liste = mitarbeiter_liste.filter(abteilung=eigene_abteilung)
    
    abteilung_filter = request.GET.get('abteilung')
    standort_filter = request.GET.get('standort')
    
    if abteilung_filter and request.user.is_staff:
        mitarbeiter_liste = mitarbeiter_liste.filter(abteilung=abteilung_filter)
    
    if standort_filter:
        mitarbeiter_liste = mitarbeiter_liste.filter(standort=standort_filter)

    context = {
        'mitarbeiter_liste': mitarbeiter_liste,
        'abteilung_filter': abteilung_filter,
        'standort_filter': standort_filter,
        'anzahl_siegburg': mitarbeiter_liste.filter(standort='siegburg').count(),
        'anzahl_bonn': mitarbeiter_liste.filter(standort='bonn').count(),
    }
    return render(request, 'arbeitszeit/mitarbeiter_uebersicht.html', context)


def register(request):
    if request.user.is_authenticated:
        return redirect('arbeitszeit:dashboard')

    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            try:
                user = User.objects.create_user(
                    username=form.cleaned_data['username'],
                    email=form.cleaned_data['email'],
                    password=form.cleaned_data['password1'],
                    first_name=form.cleaned_data['vorname'],
                    last_name=form.cleaned_data['nachname'],
                )
                
                Mitarbeiter.objects.create(
                    user=user,
                    personalnummer=form.cleaned_data['personalnummer'],
                    vorname=form.cleaned_data['vorname'],
                    nachname=form.cleaned_data['nachname'],
                    abteilung=form.cleaned_data['abteilung'],
                    standort=form.cleaned_data['standort'],
                    eintrittsdatum=form.cleaned_data['eintrittsdatum'],
                    aktiv=True
                )
                
                login(request, user)
                messages.success(request, 'Registrierung erfolgreich! Willkommen!')
                return redirect('arbeitszeit:dashboard')
                
            except Exception as e:
                messages.error(request, f'Ein Fehler ist aufgetreten: {str(e)}')
        else:
            messages.error(request, 'Bitte korrigieren Sie die Fehler im Formular.')
    else:
        form = RegisterForm()

    return render(request, 'arbeitszeit/register.html', {'form': form})


@login_required
def admin_vereinbarung_loeschen(request, pk):
    if not request.user.is_staff:
        messages.error(request, "Keine Berechtigung.")
        return redirect('arbeitszeit:dashboard')
    
    vereinbarung = get_object_or_404(Arbeitszeitvereinbarung, pk=pk)
    
    if request.method == 'POST':
        name = vereinbarung.mitarbeiter.vollname
        art = vereinbarung.get_antragsart_display()
        vereinbarung.delete()
        messages.success(request, f'Vereinbarung von {name} ({art}) wurde gelöscht.')
    
    return redirect('arbeitszeit:admin_vereinbarungen')


@login_required
def admin_vereinbarung_docx_export(request, pk):
    """Generiert ein DOCX-Dokument mit Wochen-Gruppierung."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        messages.error(request, "Das Modul 'python-docx' ist nicht installiert.")
        return redirect('arbeitszeit:admin_vereinbarung_genehmigen', pk=pk)

    if not request.user.is_staff:
        return redirect('arbeitszeit:dashboard')
    
    vereinbarung = get_object_or_404(Arbeitszeitvereinbarung, pk=pk)
    
    # 1. Laden und Sortieren
    tagesarbeitszeiten = vereinbarung.tagesarbeitszeiten.annotate(
        sort_order=_get_wochentag_sortierung()
    ).order_by('woche', 'sort_order')
    
    # 2. Gruppieren nach Wochen
    from collections import defaultdict
    wochen_daten = defaultdict(list)
    for tag in tagesarbeitszeiten:
        wochen_daten[tag.woche].append(tag)
    
    # --- DOKUMENT START ---
    doc = Document()
    
    # Layout
    section = doc.sections[0]
    section.page_height = Inches(11.69) # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Kopfbereich (Leerzeilen)
    for _ in range(4): doc.add_paragraph()
    
    # Betreff
    betreff_suffix = {
        'weiterbewilligung': ' - Weiterbewilligung',
        'verringerung': ' - Verringerung',
        'erhoehung': ' - Erhöhung',
        'beendigung': ' - Beendigung'
    }.get(vereinbarung.antragsart, '')
    
    p = doc.add_paragraph()
    run = p.add_run(f'Betreff: Flexible Arbeitszeit{betreff_suffix}')
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Stammdaten kurz auflisten
    doc.add_paragraph(f"Mitarbeiter: {vereinbarung.mitarbeiter.vollname}")
    doc.add_paragraph(f"Antragsart: {vereinbarung.get_antragsart_display()}")
    
    gueltig_bis = vereinbarung.gueltig_bis.strftime('%d.%m.%Y') if vereinbarung.gueltig_bis else 'unbefristet'
    doc.add_paragraph(f'Gültig: {vereinbarung.gueltig_ab.strftime("%d.%m.%Y")} bis {gueltig_bis}')
    
    doc.add_paragraph('\nSehr geehrte Damen und Herren,\n')
    doc.add_paragraph('hiermit wird die folgende Arbeitszeitvereinbarung getroffen:\n')
    
    # --- WOCHENVERTEILUNG GENERIEREN ---
    if wochen_daten:
        p = doc.add_paragraph()
        p.add_run('Wochenverteilung:').bold = True
        
        # Schleife über die Wochen (Woche 1, Woche 2, etc.)
        for woche in sorted(wochen_daten.keys()):
            tage = wochen_daten[woche]
            
            # Überschrift für die Woche
            p = doc.add_paragraph()
            run = p.add_run(f"Woche {woche}")
            run.bold = True
            run.font.color.rgb = RGBColor(26, 77, 46) # Dunkelgrün
            
            # Tabelle erstellen
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header Zeile
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Wochentag'
            hdr_cells[1].text = 'Arbeitszeit'
            
            # Header Styling (Grüner Hintergrund)
            for cell in hdr_cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '1a4d2e')
                cell._element.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            
            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
           # Datenzeilen
            wochen_summe = 0
            for tag in tage:
                row_cells = table.add_row().cells
                row_cells[0].text = tag.get_wochentag_display()
                # FIX: HHMM → Minuten konvertieren
                row_cells[1].text = zeitwert_to_str(hhmm_to_minuten(tag.zeitwert)) + " h"
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if tag.zeitwert:
                    # FIX: HHMM → Minuten konvertieren vor Addition
                    wochen_summe += hhmm_to_minuten(tag.zeitwert)

            # Summenzeile pro Woche hinzufügen (NACH der for-Schleife!)
            row_cells = table.add_row().cells
            row_cells[0].text = "Summe"
            row_cells[1].text = zeitwert_to_str(wochen_summe) + " h"

            # Summenzeile fett machen & Hintergrund grau
            for cell in row_cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F0F0F0') # Hellgrau
                cell._element.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph() # Abstand zur nächsten Woche
            
    # --- ENDE DOKUMENT ---

    doc.add_paragraph(f'Diese Vereinbarung gilt vom {vereinbarung.gueltig_ab.strftime("%d.%m.%Y")} bis {gueltig_bis}.')
    doc.add_paragraph('\nMit freundlichen Grüßen\n\n\n_______________________________\nUnterschrift')
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    response = HttpResponse(
        file_stream.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"Arbeitszeit_{vereinbarung.mitarbeiter.nachname}_{vereinbarung.pk}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def admin_vereinbarung_pdf_export(request, pk):
    try:
        from weasyprint import HTML
    except ImportError:
        messages.error(request, "Das Modul 'weasyprint' ist nicht installiert.")
        return redirect('arbeitszeit:admin_vereinbarung_genehmigen', pk=pk)

    if not request.user.is_staff:
        return redirect('arbeitszeit:dashboard')

    vereinbarung = get_object_or_404(Arbeitszeitvereinbarung, pk=pk)

    # Daten laden und sortieren
    tagesarbeitszeiten = vereinbarung.tagesarbeitszeiten.annotate(
        sort_order=_get_wochentag_sortierung()
    ).order_by('woche', 'sort_order')

    # DATEN GRUPPIEREN
    from collections import defaultdict
    temp_grouped = defaultdict(list)
    for tag in tagesarbeitszeiten:
        # FIX: HHMM → Minuten konvertieren
        tag.display_zeit = zeitwert_to_str(hhmm_to_minuten(tag.zeitwert))
        temp_grouped[tag.woche].append(tag)

    # Strukturierte Daten mit Summen erstellen
    wochen_daten = {}
    for woche in sorted(temp_grouped.keys()):
        tage = temp_grouped[woche]
        # FIX: HHMM → Minuten konvertieren vor Summierung
        gesamt_minuten = sum(hhmm_to_minuten(t.zeitwert) for t in tage if t.zeitwert)
        wochen_daten[woche] = {
            'tage': tage,
            'summe': zeitwert_to_str(gesamt_minuten)
        }

    context = {
        'vereinbarung': vereinbarung,
        'wochen_daten': wochen_daten,
    }

    html_string = render_to_string('arbeitszeit/pdf_vereinbarung.html', context)
    html = HTML(string=html_string, base_url=request.build_absolute_uri('/'))
    pdf = html.write_pdf()

    response = HttpResponse(pdf, content_type='application/pdf')
    filename = f"Arbeitszeit_{vereinbarung.mitarbeiter.nachname}_{vereinbarung.pk}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def mitarbeiter_detail(request, pk):
    mitarbeiter = get_object_or_404(Mitarbeiter, pk=pk)
    
    is_schichtplaner = hasattr(request.user, 'mitarbeiter') and request.user.mitarbeiter.rolle == 'schichtplaner'
    if not (request.user.is_staff or is_schichtplaner):
        messages.error(request, "Keine Berechtigung.")
        return redirect('arbeitszeit:dashboard')
    
    if request.method == 'POST':
        mitarbeiter.schichtplan_kennung = request.POST.get('schichtplan_kennung', '')
        mitarbeiter.kann_tagschicht = 'kann_tagschicht' in request.POST
        mitarbeiter.kann_nachtschicht = 'kann_nachtschicht' in request.POST
        mitarbeiter.nachtschicht_nur_wochenende = 'nachtschicht_nur_wochenende' in request.POST
        mitarbeiter.nur_zusatzdienste_wochentags = 'nur_zusatzdienste_wochentags' in request.POST
        mitarbeiter.verfuegbarkeit = request.POST.get('verfuegbarkeit', 'voll')
        mitarbeiter.schichtplan_einschraenkungen = request.POST.get('schichtplan_einschraenkungen', '')
        
        try:
            mitarbeiter.max_wochenenden_pro_monat = int(request.POST.get('max_wochenenden_pro_monat', 4))
        except ValueError:
            mitarbeiter.max_wochenenden_pro_monat = 4
        
        mitarbeiter.save()
        messages.success(request, f"Mitarbeiter {mitarbeiter.vollname} aktualisiert!")
        return redirect('schichtplan:mitarbeiter_uebersicht')
    
    return render(request, 'arbeitszeit/mitarbeiter_detail.html', {'mitarbeiter': mitarbeiter})